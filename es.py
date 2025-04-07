import streamlit as st
import base64
import time
import os
import datetime
import pandas as pd
from streamlit_extras.grid import grid
from streamlit_extras.stylable_container import stylable_container
from PIL import Image
import folium
import plotly.graph_objects as go
from geopy.geocoders import Nominatim
from openpyxl import load_workbook
import geopandas as gpd
from folium import Choropleth
from streamlit_folium import folium_static
import json
import re
from functions import *
import hashlib
import qrcode
import io
from pathlib import Path
from dotenv import load_dotenv

load_dotenv()

# Access the token
hf_token = os.getenv("HUGGINGFACE_TOKEN")


st.set_page_config(
    page_title="Blood Donation Campaign Dashboard",
    page_icon="🩸",
    layout="wide",
    initial_sidebar_state="expanded"
)
def load_data1():
        if "uploaded_data" in st.session_state:
            return st.session_state["uploaded_data"]  # Load validated data
        else:
            return pd.read_excel('last.xlsx')  # Load default file if no uploaded data

def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# Fonction pour vérifier les identifiants
def check_credentials(username, password):
    users = {
        "admin": hash_password("admin"),
    }
    if username in users and users[username] == hash_password(password):
        return True
    return False

# Fonction pour créer le fichier de utilisateurs (à exécuter une fois pour initialiser)
def create_users_file():
    if not os.path.exists("users.csv"):
        users_data = {
            "username": ["user1", "user2", "admin"],
            "password": [
                hash_password("password1"),
                hash_password("password2"),
                hash_password("admin123")
            ]
        }
        pd.DataFrame(users_data).to_csv("users.csv", index=False)
        st.success("Fichier d'utilisateurs créé avec succès!")

# Configuration de session state pour garder l'état de connexion
if "authenticated" not in st.session_state:
    st.session_state["authenticated"] = False


def login_form():
    # Chemin vers votre image locale
    img_path = "OIP.jpeg"  # Assurez-vous que le chemin est correct
    img_absolute_path = Path(img_path).absolute()
    st.image(img_absolute_path,width=500)
    st.markdown(
        f"""
        <style>
            .block-container {{
                padding-top: 3.5rem;
                padding-left: 25rem;
                padding-right: 25rem;
                position: relative;
                z-index: 1;
            }}
        </style>
        """,
        unsafe_allow_html=True,
    )

    # Interface d'authentification
    st.title("Blood Donation Dashboard ")
    st.subheader("Connect yourself  to continue")
    
    with st.form("login_form"):
        username = st.text_input("User")
        password = st.text_input("Password", type="password")
        submit = st.form_submit_button("Login")
        
        if submit:
            if check_credentials(username, password):
                st.session_state["authenticated"] = True
                st.session_state["username"] = username
                st.success("Connexion réussie! valider à nouveau pour confirmer la connexion")
                load_data1()
                st.rerun()
            else:
                st.error("Nom d'utilisateur ou mot de passe incorrect")

def get_base64_of_bin_file(bin_file):
    with open(bin_file, 'rb') as f:
        data = f.read()
    return base64.b64encode(data).decode()

def main():
    # Initialize session state
    if "authenticated" not in st.session_state:
        st.session_state["authenticated"] = False
    if "file_uploaded" not in st.session_state:
        st.session_state["file_uploaded"] = False
    if "file_validated" not in st.session_state:
        st.session_state["file_validated"] = False

    # Authentication check
    if not st.session_state["authenticated"]:
        login_form()
    elif not st.session_state["file_uploaded"]:
        file_upload_page()
    elif not st.session_state["file_validated"]:
        file_upload_page()
        st.success("File validated successfully! Proceed to the main app.")
        
    else:
        main_app()



#_____________________________________
def main_app():
    #_____________________________________________________________________________________________
    @st.cache_resource
    def get_text(key):
        if 'language' not in st.session_state:
            return key
        if 'get_text' not in st.session_state:
            return key
        return st.session_state.get_text(key)
    with st.sidebar:
        pdf_path = "Rapport_CampagneDon_Sang_Indabax.pdf"
        with open(pdf_path, "rb") as f:
            pdf_bytes = f.read()   
        st.download_button(
            label=get_text("Download the report"),
            data=pdf_bytes,
            file_name="document.pdf",
            mime="application/pdf"
        )
    #________________________
    if st.sidebar.button("Déconnexion"):
        st.session_state["authenticated"] = False
        st.rerun()

    if 'language' not in st.session_state:
        st.session_state.language = "English"

    # Chargement des traductions
    with open('langage.json', 'r', encoding='utf-8') as json_file:
        TRANSLATIONS = json.load(json_file)

    def get_modification_time(path = "donnees.xlsx"):
        return os.path.getmtime(path)

    #_____________________Function_____________________________________

    # Cache pour les shapefiles avec simplification de géométrie
    @st.cache_resource
    def load_shapefile(shapefile_path, simplify_tolerance=None):
        gdf = gpd.read_file(shapefile_path)
        if simplify_tolerance is not None:
            gdf.geometry = gdf.geometry.simplify(simplify_tolerance)
        return gdf


    # Fonction pour charger les données avec la date de modification comme hash
    @st.cache_resource
    def load_data1():
        if "uploaded_data" in st.session_state:
            return st.session_state["uploaded_data"]  # Load validated data
        else:
            return pd.read_excel('last.xlsx')  # Load default file if no uploaded data

    def load_data3():
        df = pd.read_excel('last.xlsx', sheet_name="2020")
        return df
    @st.cache_resource
    def load_data2(get_modification_time):
        df = pd.read_excel('donnees.xlsx')
        return df

    def get_combined_data():
        return pd.concat([load_data1(), load_data2(get_modification_time())], axis=0)


    #________________________


    def get_hierarchical_data():
        """Prétraiter les données avec les relations hiérarchiques entre niveaux administratifs"""
        # Charger les shapefiles avec les relations hiérarchiques
        geo_data_1 = load_shapefile("gadm41_CMR_1.shp", simplify_tolerance=0.01)
        geo_data_2 = load_shapefile("gadm41_CMR_2.shp", simplify_tolerance=0.01)
        geo_data_3 = load_shapefile("gadm41_CMR_3.shp", simplify_tolerance=0.01)
        
        # Créer un dictionnaire de correspondance entre arrondissements et départements/régions
        admin_hierarchy = pd.DataFrame({
            'NAME_3': geo_data_3['NAME_3'],
            'NAME_2': geo_data_3['NAME_2'], 
            'NAME_1': geo_data_3['NAME_1']
        })
        
        # Fusionner les données des volontaires avec la hiérarchie administrative
        combined_data = get_combined_data()
        enriched_data = combined_data.merge(
            admin_hierarchy, 
            left_on='Arrondissement_de_résidence_', 
            right_on='NAME_3', 
            how='left'
        )
        
        return enriched_data



    @st.cache_resource
    def create_card(content, key, cell_height="200px", cell_width="160px"):
        """Create a styled card with custom dimensions"""
        with stylable_container(
            key=key,
            css_styles=f"""
                {{
                    min-height: {cell_height};
                    height: {cell_height};
                    width: {cell_width};
                    max-width: {cell_width};
                    border: 1px solid #c0c0c0;
                    border-radius: 10px;
                    margin: 0px;  # Small margin for spacing
                    padding: 0px;
                    display: flex;
                    flex-direction: column;
                    align-items: center;
                    justify-content: center;
                    background-color: #FFFEBA;
                    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.2);
                    transition: all 0.2s ease;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    color: green;
                }}
                
                .card-title {{
                        font-weight: bold;
                        margin: 0px;
                        padding: 0px;
                        font-size: 1em;
                        text-align: center;
                        color: #8a2be2;  # Light purple color
                    }}
                
                .card-value {{
                    font-size:1.2em;
                    text-align: center;
                }}
                
            """
        ):
            st.markdown(content, unsafe_allow_html=True)


    def create_metric_card(column, title, plot_function, height="250px",width="100%"):
        cell_width = width
        key = f"{title.lower().replace(' ', '_')}_card"
        
        with column:
            # Create content for the card heading
            header_content = f"""
            <div class="card-title">{title}</div>
            """
            
            # Use the stylable container with reusable styling
            with stylable_container(
                key=key,
                css_styles=f"""
                    {{
                        min-height: {height};
                        height: auto;
                        width: {cell_width};
                        max-width: {cell_width};
                        border: 1px solid black;
                        border-radius: 10px;
                        background-color: #f8f9fa;
                        box-shadow: 0 4px 6px rgba(0, 0, 0, 0.8);
                        transition: all 0.2s ease;
                        overflow: hidden;
                        text-overflow: ellipsis;
                        color: #8a2be2;
                    }}
                    
                    .card-title {{
                        font-weight: bold;
                        margin-top: 0px;
                        padding: 0px;
                        font-size: 1.2em;
                        text-align: center;
                    }}
                    
                    .card-value {{
                        color: black;
                        font-size: 1em;
                        font-weight: bold;
                        text-align: center;
                    }}
                    
                """
            ):
                # Render the header content
                st.markdown(header_content, unsafe_allow_html=True)
                
                # Get the plot from the provided function
                fig = plot_function()
            




    def get_preprocessed_data(level):
        """Prétraiter les données pour différents niveaux administratifs"""
        enriched_data = get_hierarchical_data()
        
        if level == 1:  # Région
            data_counts = enriched_data['NAME_1'].value_counts().reset_index()
            data_counts.columns = ['NAME_1', 'Nb']
            return data_counts
            
        elif level == 2:  # Département
            data_counts = enriched_data['NAME_2'].value_counts().reset_index()
            data_counts.columns = ['NAME_2', 'Nb']
            return data_counts
            
        elif level == 3:  # Arrondissement
            data_counts = enriched_data['Arrondissement_de_résidence_'].value_counts().reset_index()
            data_counts.columns = ['Arrondissement_de_résidence_', 'Nb']
            return data_counts
        
        return None



    # Helper function to format card content
    @st.cache_resource
    def format_card_content(title, value):
        
        return f"""
            <div class='card-title' style='margin: 0px; padding: 0px;'>{title}</div>
            <div class='card-value' style='margin: 0px; padding: 0px;'>{value}</div>
        """
    #<div class='card-delta' style='margin: 0px; padding: 0px;'>{delta_text}</div>______________________________________________________________

    if "donors" not in st.session_state:
        st.session_state.donors = pd.DataFrame(columns=[
        'id', 'age', 'Date_remplissage','Date_naiss', 'niveau_detude', 'genre', 'taille', 'poids',
        'situation_matrimoniale', 'profession', 'arrondissement_residence',
        'quartier_residence', 'nationalite', 'religion', 'deja_donne_sang',
        'date_dernier_don', 'taux_dhemoglobine', 'eligibilite_au_don',
        'est_sous_anti_biotherapie', 'taux_dhemoglobine_bas', 'date_dernier_don_3_mois',
        'ist_recente', 'ddr_incorrecte', 'allaitement', 'accouchement_6mois',
        'interruption_grossesse', 'enceinte', 'antecedent_transfusion',
        'porteur_hiv_hbs_hcv', 'opere', 'drepanocytaire', 'diabetique',
        'hypertendus', 'asthmatiques', 'cardiaque', 'tatoue', 'scarifie',
        'autres_raisons'
    ])
            
            # Add a sample donor record
    # Import the challenge dataset
    #challenge_data = pd.read_excel('Challenge dataset.xlsx')
    # Apply custom CSS for the sidebar styling

    st.markdown('''
    <style>
        /* Sidebar styles */
        .sidebar , .sidebar-content {
            display: flex;
            flex-direction: column;
            align-items: flex-start;
            justify-content: flex-start;
        }

        .sidebar, .styled-header {
            font-size: 18px;
            font-weight: bold;
            color: WHITE !important;
            text-align: center;
            margin-top: 20px;
            background-color: rgba(25, 0, 0, 0.4);
            padding: 10px;
            border-radius: 6px;
            box-shadow: 10 4px 8px rgba(0, 0, 0, 0.1);
        }

        /* Styling for the sidebar links */
        .sidebar, .sidebar-nav {
            display: flex;
            flex-direction: column;
            align-items: flex-start;
            width: 100%;
        }

        .sidebar-nav a {
            text-decoration: none;
            padding: 12px;
            color: blue;
            font-weight: bold;
            margin-top: 10px;
            width: 100%;
            text-align: left;
            border-radius: 5px;
            background-color: #f1f1f1;
        }

        .sidebar-nav a:hover {
            background-color: #e1e1e1;
        }
        .stButton{
            color: white;  /* Nouvelle couleur du texte lors du survol */
            font-weight: bold;  /* Nouveau poids de la police lors du survol */
            
        }
        .stButton>button:hover {
            background-color:rgb(90, 6, 6);  /* Nouvelle couleur de fond lors du survol */
            color: white;  /* Nouvelle couleur du texte lors du survol */
            font-weight: bold;
            transition: background-color 0.3s 
        }

        }
    </style>
    ''', unsafe_allow_html=True)
    st.sidebar.image('logo.jpeg', width=200)

    # Define the menu items

    def get_text(key):
        if key not in TRANSLATIONS[st.session_state.language]:
            return key  # Return key if translation doesn't exist
        return TRANSLATIONS[st.session_state.language][key]

    menu_items = {
        "Home": f"🏠 {get_text('Home')}",
        "Donations": f"💉 {get_text('Donations')}",
        "Cartography": f"📍 {get_text('Cartography')}",
        "Dataset Insights": f"🔄 {get_text('Dataset Insights')}",
        "Eligibility and Profile": f"🩺 {get_text('Eligibility and Profile')}",
        "Campaign Insights": f"📊 {get_text('Campaign Insights')}",
        "Fidélisation": f"📊 {get_text('Fidélisation')}",
        "Donor Insights": f"📊 {get_text('Donor Insights')}",
        "AI Extensions": f"🤖 {get_text('AI Extensions')}",
        "Options": f"⚙️ {get_text('Options')}",
        "About": f"ℹ️ {get_text('About')}"
    }

    # Create a selectbox for the menu
    selected_item = st.sidebar.selectbox(
        " ",
        options=list(menu_items.keys()), 
        format_func=lambda x: menu_items[x],
        index=0,  # Premier élément sélectionné par défaut
        key="main_menu",
        help="Sélectionnez une option du menu"
    )



    if selected_item =="Donor Insights":
                
        # CSS amélioré avec plus de rouge
        st.markdown("""
            <style>
            .main {
                background-color: #f8f9fa; /* Blanc cassé élégant */
                padding: 20px;
            }
            h1, h2 {
                color: #c0392b; /* Rouge sang pour les titres */
                font-family: 'Helvetica Neue', sans-serif;
                font-weight: bold;
                font-size: 32px;
                text-align: center;
            }
            .stMarkdown {
                font-family: 'Helvetica Neue', sans-serif;
                color: #333333; /* Gris anthracite doux */
            }
            .metric-card {
                border-radius: 10px;
                padding: 25px;
                margin: 15px 0;
                text-align: center;
                background-color: #ffffff;
                box-shadow: 0 6px 12px rgba(0, 0, 0, 0.1);
                transition: all 0.3s ease;
                border: 1px solid #e0e0e0;
            }
            .metric-card:hover {
                transform: translateY(-8px);
                box-shadow: 0 12px 24px rgba(0, 0, 0, 0.15);
                border: 1px solid #c0392b; /* Rouge au survol */
            }
            .metric-title {
                font-size: 22px;
                font-weight: bold;
                margin-bottom: 12px;
                color: #333333;
            }
            .metric-value {
                font-size: 32px;
                font-weight: bold;
                color: #ffffff;
                border-radius: 8px;
                padding: 8px 16px;
                display: inline-block;
            }
            </style>
        """, unsafe_allow_html=True)

        # Lecture des données
        df =load_data3()
        df['Age'] = pd.to_numeric(df['Age'], errors='coerce')
        df = df.dropna(subset=['Age'])
        df['Age'] = df['Age'].astype(int)
        df['Type de donation'] = df['Type de donation'].fillna('Sang total')
        df['Type de donation'] = df['Type de donation'].replace({'F': 'Sang total', 'B': 'Sang total'})

        # Partie 2 : Indicateurs clés
        st.header("Indicateurs clés")

        # Calcul des indicateurs
        total_donors = len(df)
        total_dons_max = (df[df['Sexe'] == 'M']['Age'].count() * 6) + (df[df['Sexe'] == 'F']['Age'].count() * 3)
        donation_goal = 10000
        percent_achieved = (total_donors / donation_goal) * 100
        most_donated_blood = df['Groupe Sanguin ABO / Rhesus'].value_counts().idxmax()
        most_frequent_don_type = df['Type de donation'].value_counts().idxmax()
        #average_age = df['Age'].mean()

        # Palette de couleurs avec plus de rouge
        professional_colors = {
            'total_donors': '#ffe6e6',      # Rouge pâle
            'percent_achieved': '#ffcccc',  # Rouge très clair
            'most_donated_blood': '#e6f0fa',# Bleu clair (contraste)
            'most_frequent_don_type': '#f5f5f5', # Gris clair
            'average_age': '#e8f5e9'        # Vert pâle
        }

        # Grille de 5 colonnes pour les indicateurs
        cols = st.columns(4)

        with cols[0]:
            st.markdown(
                f"""
                <div class="metric-card" style="background-color: {professional_colors['total_donors']};">
                    <div class="metric-title">Nombre total de donneurs</div>
                    <div class="metric-value" style="background: linear-gradient(90deg, #c0392b 0%, #8e1e18 100%);">{total_donors}</div>
                </div>
                """, unsafe_allow_html=True
            )

        with cols[1]:
            st.markdown(
                f"""
                <div class="metric-card" style="background-color: {professional_colors['percent_achieved']};">
                    <div class="metric-title">Pourcentage atteint</div>
                    <div class="metric-value" style="background: linear-gradient(90deg, #e74c3c 0%, #a93226 100%);">{percent_achieved:.2f}%</div>
                </div>
                """, unsafe_allow_html=True
            )

        with cols[2]:
            st.markdown(
                f"""
                <div class="metric-card" style="background-color: {professional_colors['most_donated_blood']};">
                    <div class="metric-title">Sang le plus donné</div>
                    <div class="metric-value" style="background: linear-gradient(90deg, #2980b9 0%, #1b5678 100%);">{most_donated_blood}</div>
                </div>
                """, unsafe_allow_html=True
            )

        with cols[3]:
            st.markdown(
                f"""
                <div class="metric-card" style="background-color: {professional_colors['most_frequent_don_type']};">
                    <div class="metric-title">Type de don fréquent</div>
                    <div class="metric-value" style="background: linear-gradient(90deg, #7f8c8d 0%, #566061 100%);">{most_frequent_don_type}</div>
                </div>
                """, unsafe_allow_html=True
            )

        # with cols[4]:
        #     st.markdown(
        #         f"""
        #         <div class="metric-card" style="background-color: {professional_colors['average_age']};">
        #             <div class="metric-title">Âge moyen</div>
        #             <div class="metric-value" style="background: linear-gradient(90deg, #27ae60 0%, #1b7944 100%);">{average_age:.1f} ans</div>
        #         </div>
        #         """, unsafe_allow_html=True
        #     )

        # Partie 3 : Analyse des donneurs
        horodateur(df)
        row1 = st.columns(2)
        row2 = st.columns(2)
        row3 = st.columns(1)
    
        st.header("Analyse des donneurs")
        # Graphique 1 : Répartition par groupes sanguins (Donut chart)
        with row1[0]:
            blood_group_counts = df['Groupe Sanguin ABO / Rhesus'].value_counts()
            fig1 = go.Figure(data=[
                go.Pie(labels=blood_group_counts.index, values=blood_group_counts.values, 
                    textinfo='label+percent', insidetextorientation='radial',
                    marker=dict(colors=['#c0392b', '#e74c3c', '#3498db', '#7f8c8d', '#2980b9', '#e6f0fa', '#f0f4f8', '#ffcccc']),
                    hole=0.4)
            ])
            fig1.update_layout(
                title="Répartition par groupes sanguins",
                title_font_size=18,
                title_x=0.5,
                showlegend=True,
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(t=50, b=50, l=50, r=50)
            )
            st.plotly_chart(fig1, use_container_width=True)

        # Graphique 2 : Répartition par âges (Histogramme)
        with row1[1]:
            fig2 = px.histogram(df, x='Age', nbins=15, color_discrete_sequence=['#c0392b'],
                                labels={'Age': 'Âge', 'count': 'Nombre de donneurs'})
            fig2.add_vline(x=df['Age'].mean(), line_dash="dash", line_color="#e74c3c",
                        annotation_text=f"Moyenne: {df['Age'].mean():.1f}", annotation_position="top right")
            fig2.update_layout(
                title="Répartition par âges",
                title_font_size=18,
                title_x=0.5,
                xaxis_title="Âge",
                yaxis_title="Nombre de donneurs",
                bargap=0.1,
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(t=50, b=50, l=50, r=50)
            )
            st.plotly_chart(fig2, use_container_width=True)

        # Graphique 3 : Sexe par groupe sanguin (Barres empilées)
        with row2[0]:
            sex_blood_group = pd.crosstab(df['Groupe Sanguin ABO / Rhesus'], df['Sexe'])
            fig3 = go.Figure(data=[
                go.Bar(name='Hommes', x=sex_blood_group.index, y=sex_blood_group['M'], marker_color='#c0392b'),  # Rouge naturel
                go.Bar(name='Femmes', x=sex_blood_group.index, y=sex_blood_group['F'], marker_color='#3498db')   # Bleu naturel
            ])
            fig3.update_layout(
                barmode='stack',
                title="Sexe par groupe sanguin",
                title_font_size=18,
                title_x=0.5,
                xaxis_title="Groupe sanguin",
                yaxis_title="Nombre de donneurs",
                legend_title="Sexe",
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(t=50, b=50, l=50, r=50)
            )
            st.plotly_chart(fig3, use_container_width=True)

        # Graphique 4 : Âge par groupe sanguin (Box plot)
        with row2[1]:
            fig4 = go.Figure()
            for i, blood_group in enumerate(df['Groupe Sanguin ABO / Rhesus'].unique()):
                fig4.add_trace(go.Box(
                    y=df[df['Groupe Sanguin ABO / Rhesus'] == blood_group]['Age'],
                    name=blood_group,
                    marker_color=['#c0392b', '#e74c3c', '#3498db', '#7f8c8d', '#2980b9', '#ffcccc', '#f0f4f8', '#e6f0fa'][i % 8],
                    boxpoints='outliers',
                    jitter=0.3,
                    line_width=1.5
                ))
            fig4.update_layout(
                title="Âge par groupe sanguin",
                title_font_size=18,
                title_x=0.5,
                xaxis_title="Groupe sanguin",
                yaxis_title="Âge",
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(t=50, b=50, l=50, r=50),
                showlegend=True
            )
            st.plotly_chart(fig4, use_container_width=True)

        # Graphique 5 : Lien entre groupes sanguins et types de don en pourcentage (Barres empilées)
        with row3[0]:
            blood_group_donation = pd.crosstab(df['Groupe Sanguin ABO / Rhesus'], df['Type de donation'])
            blood_group_donation_pct = blood_group_donation.div(blood_group_donation.sum(axis=1), axis=0) * 100
            fig5_new = go.Figure()
            colors = ['#c0392b', '#3498db', '#7f8c8d']  # Rouge, Bleu, Gris naturel
            for i, donation_type in enumerate(blood_group_donation_pct.columns):
                fig5_new.add_trace(go.Bar(
                    x=blood_group_donation_pct.index,
                    y=blood_group_donation_pct[donation_type],
                    name=donation_type,
                    marker_color=colors[i % len(colors)],
                    text=[f"{val:.1f}%" for val in blood_group_donation_pct[donation_type]],
                    textposition='inside'
                ))
            fig5_new.update_layout(
                barmode='stack',
                title="Pourcentage des types de don par groupe sanguin",
                title_font_size=18,
                title_x=0.5,
                xaxis_title="Groupe sanguin",
                yaxis_title="Pourcentage (%)",
                legend_title="Type de donation",
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(t=50, b=50, l=50, r=50),
                yaxis=dict(range=[0, 100])
            )
            st.plotly_chart(fig5_new, use_container_width=True)

        # Graphique 6 : Répartition par type de donation (Barres empilées)
        st.subheader("Répartition par type de donation")
        donation_type_sex = pd.crosstab(df['Type de donation'], df['Sexe'])
        fig6 = go.Figure(data=[
            go.Bar(name='Hommes', x=donation_type_sex.index, y=donation_type_sex['M'], marker_color='#c0392b'),  # Rouge naturel
            go.Bar(name='Femmes', x=donation_type_sex.index, y=donation_type_sex['F'], marker_color='#3498db')   # Bleu naturel
        ])
        fig6.update_layout(
            barmode='stack',
            title="Répartition par type de donation",
            title_font_size=18,
            title_x=0.5,
            xaxis_title="Type de donation",
            yaxis_title="Nombre de donneurs",
            legend_title="Sexe",
            height=400,
            paper_bgcolor='rgba(0,0,0,0)',
            plot_bgcolor='rgba(0,0,0,0)',
            margin=dict(t=50, b=50, l=50, r=50)
        )
        st.plotly_chart(fig6, use_container_width=True)

        # Partie : Atteintes des objectifs par groupe sanguin
        st.header("Atteintes des objectifs par groupe sanguin")

        # Objectif total et répartition basée sur la demande globale
        donation_goal = 10000
        blood_group_demand = {
            'O+': 0.35, 'A+': 0.30, 'B+': 0.10, 'AB+': 0.05,
            'O-': 0.07, 'A-': 0.06, 'B-': 0.02, 'AB-': 0.01
        }
        blood_groups = df['Groupe Sanguin ABO / Rhesus'].unique()
        blood_group_counts = df['Groupe Sanguin ABO / Rhesus'].value_counts().to_dict()

        # Calcul des objectifs par groupe sanguin
        goal_per_group = {bg: blood_group_demand.get(bg, 0.05) * donation_goal for bg in blood_groups}
        collected_per_group = blood_group_counts

        # Filtre interactif
        selected_blood_group = st.selectbox("Choisir un groupe sanguin", options=['Tous'] + list(blood_groups))

        # Disposition en 2 colonnes
        obj_cols = st.columns([2, 1])

        # Graphique 7 : Comparaison collecté vs objectif (Barres juxtaposées)
        with obj_cols[0]:
            if selected_blood_group == 'Tous':
                fig7 = go.Figure(data=[
                    go.Bar(name='Objectif', x=list(goal_per_group.keys()), y=list(goal_per_group.values()), marker_color='#3498db'),  # Bleu naturel
                    go.Bar(name='Collecté', x=list(goal_per_group.keys()), y=[collected_per_group.get(bg, 0) for bg in goal_per_group.keys()], marker_color='#c0392b')  # Rouge naturel
                ])
            else:
                fig7 = go.Figure(data=[
                    go.Bar(name='Objectif', x=[selected_blood_group], y=[goal_per_group[selected_blood_group]], marker_color='#3498db'),  # Bleu naturel
                    go.Bar(name='Collecté', x=[selected_blood_group], y=[collected_per_group.get(selected_blood_group, 0)], marker_color='#c0392b')  # Rouge naturel
                ])
            fig7.update_layout(
                barmode='group',
                title=f"Collecté vs Objectif ({selected_blood_group})",
                title_font_size=18,
                title_x=0.5,
                xaxis_title="Groupe sanguin",
                yaxis_title="Nombre de poches",
                height=400,
                paper_bgcolor='rgba(0,0,0,0)',
                plot_bgcolor='rgba(0,0,0,0)',
                margin=dict(t=50, b=50, l=50, r=50),
                showlegend=True
            )
            st.plotly_chart(fig7, use_container_width=True)

        # Graphique 8 : Jauge (Gauge) pour le pourcentage atteint
        with obj_cols[1]:
            if selected_blood_group == 'Tous':
                total_collected = sum(collected_per_group.values())
                total_goal = donation_goal
                percent_achieved_total = (total_collected / total_goal) * 100
                options8 = {
                    "series": [
                        {
                            "type": "gauge",
                            "startAngle": 180,
                            "endAngle": 0,
                            "min": 0,
                            "max": 100,
                            "splitNumber": 10,
                            "itemStyle": {"color": "#c0392b"},
                            "progress": {"show": True, "width": 20, "itemStyle": {"color": "#e74c3c"}},
                            "pointer": {"show": True},
                            "axisLine": {"lineStyle": {"width": 20, "color": [[1, "#ffcccc"]]}},
                            "axisTick": {"show": False},
                            "splitLine": {"length": 15, "lineStyle": {"width": 2, "color": "#999"}},
                            "axisLabel": {"distance": 25, "color": "#999", "fontSize": 12},
                            "detail": {
                                "valueAnimation": True,
                                "fontSize": 20,
                                "color": "#c0392b",
                                "offsetCenter": [0, "70%"],
                                "formatter": "{value}%"
                            },
                            "data": [{"value": percent_achieved_total, "name": "Atteint"}]
                        }
                    ]
                }
                st_echarts(options=options8, height="400px", key="gauge_total")
            else:
                collected = collected_per_group.get(selected_blood_group, 0)
                goal = goal_per_group[selected_blood_group]
                percent_achieved_group = (collected / goal) * 100
                options8 = {
                    "series": [
                        {
                            "type": "gauge",
                            "startAngle": 180,
                            "endAngle": 0,
                            "min": 0,
                            "max": 100,
                            "splitNumber": 10,
                            "itemStyle": {"color": "#c0392b"},
                            "progress": {"show": True, "width": 20, "itemStyle": {"color": "#e74c3c"}},
                            "pointer": {"show": True},
                            "axisLine": {"lineStyle": {"width": 20, "color": [[1, "#ffcccc"]]}},
                            "axisTick": {"show": False},
                            "splitLine": {"length": 15, "lineStyle": {"width": 2, "color": "#999"}},
                            "axisLabel": {"distance": 25, "color": "#999", "fontSize": 12},
                            "detail": {
                                "valueAnimation": True,
                                "fontSize": 20,
                                "color": "#c0392b",
                                "offsetCenter": [0, "70%"],
                                "formatter": "{value}%"
                            },
                            "data": [{"value": round(percent_achieved_group, 2), "name": f"Atteint ({selected_blood_group})"}]
                        }
                    ]
                }
                st_echarts(options=options8, height="400px", key=f"gauge_{selected_blood_group}")
    elif selected_item =="Fidélisation":
        st.markdown("""
            <style>
            .block-container {
                padding-top: 0rem;
                padding-bottom: 0rem;
                padding-left:4rem;
                padding-right: 4rem;
                margin-top: 0px;
            }
                
                /* Remove extra padding around header */
                header {
                    margin-bottom: 0rem !important;
                    padding-bottom: 0rem !important;
                }
            </style>
            """, unsafe_allow_html=True)
        
        combined_data = get_combined_data()
        row1_cols = st.columns(3)
        with row1_cols[1]:
            all_arrondissement = combined_data['Arrondissement_de_résidence_'].unique()
            selected_arrondissement = st.sidebar.multiselect(
                get_text("Districts"),
                all_arrondissement
            )
            if not selected_arrondissement:
                selected_arrondissement = all_arrondissement
        with row1_cols[2]:
            all_case = combined_data['ÉLIGIBILITÉ_AU_DON.'].unique()
            selected_case = st.sidebar.multiselect(
                "Eligible",
                all_case # Limiter par défaut pour améliorer les performances
            )
            if not selected_case:
                selected_case = all_case
        # Filtrer les données avec tous les filtres
        filtered_data = combined_data[
            (combined_data['Arrondissement_de_résidence_'].isin(selected_arrondissement)) &
            (combined_data['ÉLIGIBILITÉ_AU_DON.'].isin(selected_case))]
        
    # Étape 1 : Filtrer les donneurs éligibles
        df = load_data1()
        df_eligible = df[df['ÉLIGIBILITÉ_AU_DON.'] == 'Eligible']

        # Identifier les donneurs récurrents et non récurrents
        df_eligible_recurrent = df_eligible[df_eligible['A-t-il (elle) déjà donné le sang'] == 'Oui']
        df_eligible_non_recurrent = df_eligible[df_eligible['A-t-il (elle) déjà donné le sang'] == 'Non']

        # Liste des colonnes démographiques à analyser
        demographic_columns = ['Classe_Age', 'Genre_', "Niveau_d'etude", 'Religion_Catégorie', 
                            'Situation_Matrimoniale_(SM)', 'categories', 'Arrondissement_de_résidence_']

        # Fonction pour générer un graphique interactif avec les top 4 catégories
        def plot_top4_demographic(data_recurrent, data_non_recurrent, column, title_prefix, comparison=False, orientation='v'):
            count_recurrent = data_recurrent[column].value_counts()
            if comparison:
                count_non_recurrent = data_non_recurrent[column].value_counts()
                
                all_categories = pd.concat([count_recurrent, count_non_recurrent], axis=1, sort=False)
                all_categories.columns = ['Récurrents', 'Non Récurrents']
                all_categories.fillna(0, inplace=True)
                
                all_categories['Total'] = all_categories['Récurrents'] + all_categories['Non Récurrents']
                top4_categories = all_categories.sort_values('Total', ascending=False).head(4).index
                
                count_recurrent = count_recurrent[count_recurrent.index.isin(top4_categories)]
                count_non_recurrent = count_non_recurrent[count_non_recurrent.index.isin(top4_categories)]
                
                fig = go.Figure()
                fig.add_trace(go.Bar(
                    x=count_recurrent.index if orientation == 'v' else count_recurrent.values,
                    y=count_recurrent.values if orientation == 'v' else count_recurrent.index,
                    name='Récurrents (Oui)',
                    marker_color='#EC8282',
                    text=count_recurrent.values,
                    textposition='auto',
                    orientation='h' if orientation == 'h' else 'v'
                ))
                fig.add_trace(go.Bar(
                    x=count_non_recurrent.index if orientation == 'v' else count_non_recurrent.values,
                    y=count_non_recurrent.values if orientation == 'v' else count_non_recurrent.index,
                    name='Non Récurrents (Non)',
                    marker_color='#ff5733',
                    text=count_non_recurrent.values,
                    textposition='auto',
                    orientation='h' if orientation == 'h' else 'v'
                ))
            else:
                top4_categories = count_recurrent.head(4).index
                count_recurrent = count_recurrent[count_recurrent.index.isin(top4_categories)]
                
                fig = go.Figure()
                fig.add_trace(go.Bar(
                    x=count_recurrent.index if orientation == 'v' else count_recurrent.values,
                    y=count_recurrent.values if orientation == 'v' else count_recurrent.index,
                    name='Récurrents',
                    marker_color='#EC8282',
                    text=count_recurrent.values,
                    textposition='auto',
                    orientation='h' if orientation == 'h' else 'v'
                ))
            
            fig.update_layout(
                title=f"{title_prefix} par {column} (Top 4)",
                xaxis_title=column if orientation == 'v' else 'Nombre de donneurs',
                yaxis_title='Nombre de donneurs' if orientation == 'v' else column,
                xaxis=dict(tickangle=45),
                legend=dict(title='Statut de récurrence', orientation='h', yanchor='bottom', y=-0.3, xanchor='center', x=0.5),
                plot_bgcolor='white',
                width=800,
                height=600,
                barmode='group' if comparison else 'stack'
            )
            
            return fig

        # Initialize Streamlit
        st.title(get_text("Blood Donor Analysis"))
        st.write("Select a demographic variable to visualize the distribution of recurrent and non-recurrent donors.")

        # Sélection des paramètres par l'utilisateur
        selected_column = st.selectbox(get_text("Choose a demographic variable:"), demographic_columns)
        orientation = st.radio(get_text("Graph orientation:"), ['Vertical', 'Horizontal'])
        comparison = st.checkbox(get_text("Compare recurrent and non-recurrent donors?"), value=False)

        # Convertir en format utilisable par la fonction
        orientation_value = 'v' if orientation == 'Verticale' else 'h'

        # Générer le graphique
        fig = plot_top4_demographic(df_eligible_recurrent, df_eligible_non_recurrent, selected_column, "Analyse des donneurs", comparison, orientation_value)

        # Afficher le graphique dans Streamlit
        st.plotly_chart(fig)

        ###########################"
        df_temp_non_eligible = df[df['ÉLIGIBILITÉ_AU_DON.'] == 'Temporairement Non-eligible']

        # Séparer les hommes et les femmes
        df_temp_men = df_temp_non_eligible[df_temp_non_eligible['Genre_'] == 'Homme']
        df_temp_women = df_temp_non_eligible[df_temp_non_eligible['Genre_'] == 'Femme']

        # Définir une palette de couleurs
        color_men = '#F50307'
        color_women = '#ff7f0e'

        # Fonction pour générer un graphique Plotly
        def plot_top4_demographic(data_men, data_women, column, title_prefix, orientation='v'):
            count_men = data_men[column].value_counts()
            count_women = data_women[column].value_counts()
            
            all_categories = pd.concat([count_men, count_women], axis=1, sort=False).fillna(0)
            all_categories.columns = ['Hommes', 'Femmes']
            all_categories['Total'] = all_categories['Hommes'] + all_categories['Femmes']
            top4_categories = all_categories.sort_values('Total', ascending=False).head(4).index
            
            count_men = count_men[count_men.index.isin(top4_categories)]
            count_women = count_women[count_women.index.isin(top4_categories)]
            
            fig = go.Figure()
            
            if orientation == 'v':
                fig.add_trace(go.Bar(x=count_men.index, y=count_men.values, name='Hommes', marker_color=color_men))
                fig.add_trace(go.Bar(x=count_women.index, y=count_women.values, name='Femmes', marker_color=color_women))
                fig.update_layout(xaxis_title=column, yaxis_title='Donor number')
            else:
                fig.add_trace(go.Bar(y=count_men.index, x=count_men.values, name='Hommes', marker_color=color_men, orientation='h'))
                fig.add_trace(go.Bar(y=count_women.index, x=count_women.values, name='Femmes', marker_color=color_women, orientation='h'))
                fig.update_layout(xaxis_title='Donor', yaxis_title=column)
            
            return fig
            # Interface Streamlit
        st.title(get_text("Analysis of Temporarily Non-Eligible Donors"))

        demographic_columns = {
            'Classe_Age': get_text('Age Group'),
            'categories': get_text('Professional Categories'),
            'Arrondissement_de_résidence_': get_text('District of Residence'),
            'Raison_indisponibilité_fusionnée': get_text('Reasons for Ineligibility')
        }

        selected_column = st.selectbox(get_text("select a catégory"), list(demographic_columns.keys()), format_func=lambda x: demographic_columns[x])

        graph_orientation = 'h' if selected_column in ['categories', 'Arrondissement_de_résidence_', 'Raison_indisponibilité_fusionnée'] else 'v'

        st.plotly_chart(plot_top4_demographic(df_temp_men, df_temp_women, selected_column, "Profil des donneurs", orientation=graph_orientation))
    #_________________________________________________________________
        def plot_bar_reasons(df, eligibility_type='Temporairement Non-eligible', gender=None):
            filtered_df = df[df['ÉLIGIBILITÉ_AU_DON.'] == eligibility_type]
            if gender:
                filtered_df = filtered_df[filtered_df['Genre_'] == gender]
            
            reasons_list = filtered_df['Raison_indisponibilité_fusionnée'].dropna().str.split(';').explode().str.strip()
            reason_counts = Counter(reasons_list)
            
            fig = go.Figure()
            fig.add_trace(go.Bar(
                x=list(reason_counts.keys()),
                y=list(reason_counts.values()),
                marker_color="#F50307",
                text=[str(val) for val in reason_counts.values()],
                textposition='auto',
                textfont=dict(size=16, color='white', family='Arial Black')
            ))
            
            title = f"{get_text('Reasons for ineligibility')} ({eligibility_type})"
            if gender:
                title += f" - {gender}"
            
            fig.update_layout(
                title=title,
                xaxis_title=get_text("Ineligibility reasons"),
                yaxis_title=get_text("Number of occurrences"),
                template='plotly_white',
                bargap=0.2,
                xaxis=dict(tickangle=-45)
            )
            
            return fig

        def plot_frequencies_by_category(df, category_col, gender_col):
            count_data = df.groupby([category_col, gender_col]).size().unstack(fill_value=0)
            categories = count_data.index
            
            fig = go.Figure()
            for gender in count_data.columns:
                fig.add_trace(go.Bar(
                    x=categories, 
                    y=count_data[gender], 
                    name=gender,
                ))
            
            fig.update_layout(
                title=f"{get_text('Distribution of')} {category_col} {get_text('by gender')}",
                xaxis_title=category_col,
                yaxis_title=get_text("Number of people"),
                barmode='group'
            )
            return fig

        def plot_hemoglobin_box(df, eligibility_col='ÉLIGIBILITÉ_AU_DON.', hemoglobin_col='Taux d’hémoglobine'):
            if not all(col in df.columns for col in [eligibility_col, hemoglobin_col]):
                st.error(get_text("The required columns are missing from the dataset."))
                return None
            
            df[hemoglobin_col] = pd.to_numeric(df[hemoglobin_col], errors='coerce')
            df_clean = df.dropna(subset=[hemoglobin_col])
            
            fig = go.Figure()
            colors = ['#FF4040', '#FF8C00', '#FFB6C1']
            
            for i, status in enumerate(df_clean[eligibility_col].unique()):
                status_data = df_clean[df_clean[eligibility_col] == status][hemoglobin_col]
                fig.add_trace(go.Box(
                    y=status_data,
                    name=status,
                    marker_color=colors[i % len(colors)],
                    boxpoints='all', jitter=0.3, pointpos=-1.8, line_width=2
                ))
            
            fig.update_layout(
                title=get_text("Hemoglobin level distribution by eligibility status"),
                xaxis_title=get_text("Eligibility status"),
                yaxis_title=get_text("Hemoglobin level (g/dL)"),
                template='plotly_white',
                showlegend=True,
                height=600
            )
            
            fig.add_hline(y=12.5, line_dash="dash", line_color="red", annotation_text=get_text("Min threshold (F)"), annotation_position="top right")
            fig.add_hline(y=13.0, line_dash="dash", line_color="blue", annotation_text=get_text("Min threshold (H)"), annotation_position="top right")
            
            return fig

        st.title(get_text("Donor Ineligibility Analysis"))
        col = st.columns(2)
        with col[0]:
            type_eligibility = st.selectbox(get_text("Choose an eligibility type"), ['Temporairement Non-eligible', 'Définitivement non-eligible'])
        with col[1]:
            gender = st.selectbox(get_text("Filter by gender (optional)"), [get_text('All'), 'Homme', 'Femme'])
        if gender == get_text('All'):
            gender = None
        st.plotly_chart(plot_bar_reasons(df, type_eligibility, gender))

        demographic_columns = {
            'Classe_Age': get_text("Age range"),
            'categories': get_text("Professional categories"),
            'Arrondissement_de_résidence_': get_text("Residence district"),
            'Raison_indisponibilité_fusionnée': get_text("Ineligibility reasons")
        }
        selected_column = st.selectbox(get_text("Select a category to analyze"), list(demographic_columns.keys()), format_func=lambda x: demographic_columns[x])
        graph_orientation = 'h' if selected_column in ['categories', 'Arrondissement_de_résidence_', 'Raison_indisponibilité_fusionnée'] else 'v'
        st.plotly_chart(plot_frequencies_by_category(df_temp_non_eligible, selected_column, 'Genre_'))
        st.title(get_text("Hemoglobin Level Analysis"))
        st.plotly_chart(plot_hemoglobin_box(df))
        #_______________________________________________________

    elif selected_item =="Campaign Insights":
        st.markdown("""
    <style>
        .card-title {
            font-weight: bold;
            font-size: 1.2em;
            text-align: center;
            padding: 10px 0;
            color: #8a2be2;
        }
        
        /* Make plots use full width of their containers */
        .js-plotly-plot, .plotly, .plot-container {
            width: 100% !important;
        }
        
        /* Remove default padding from columns to maximize space */
        div[data-testid="column"] {
            padding: 0 !important;
            margin: 0 !important;
        }
        
        /* Style for the streamlit columns inside the container */
        div[data-testid="stHorizontalBlock"] {
            width: 100% !important;
            gap: 5px;
        }
    </style>
    """, unsafe_allow_html=True)
        st.markdown(
        """
        <style>
            div[data-baseweb="select"] > div {
                min-height: 10px !important;  /* Adjust height */
            }
        </style>
        """,
        unsafe_allow_html=True
    )
        st.markdown("""
            <style>
            .block-container {
                padding-top: 0rem;
                padding-bottom: 0rem;
                padding-left:4rem;
                padding-right: 4rem;
                margin-top: 0px;
            }
                
                /* Remove extra padding around header */
                header {
                    margin-bottom: 0rem !important;
                    padding-bottom: 0rem !important;
                }
            </style>
            """, unsafe_allow_html=True)
        combined_data = get_combined_data()
        row1_cols = st.columns(3)
        with row1_cols[1]:
            all_arrondissement = combined_data['Arrondissement_de_résidence_'].unique()
            selected_arrondissement = st.sidebar.multiselect(
                get_text("Districts"),
                all_arrondissement
            )
            if not selected_arrondissement:
                selected_arrondissement = all_arrondissement
        with row1_cols[2]:
            all_case = combined_data['ÉLIGIBILITÉ_AU_DON.'].unique()
            selected_case = st.sidebar.multiselect(
                "Eligible",
                all_case # Limiter par défaut pour améliorer les performances
            )
            if not selected_case:
                selected_case = all_case
        # Filtrer les données avec tous les filtres
        filtered_data = combined_data[
            (combined_data['Arrondissement_de_résidence_'].isin(selected_arrondissement)) &
            (combined_data['ÉLIGIBILITÉ_AU_DON.'].isin(selected_case))]
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("Welcome to the  campaign Insight Section")}
        </div> """,  unsafe_allow_html=True)
        filtered_data['Date de remplissage de la fiche'] = pd.to_datetime(load_data1()['Date de remplissage de la fiche'])
        m= filtered_data['Date de remplissage de la fiche'].dt.month
        month_counts = m.value_counts().sort_index()
        categories = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
        data = {"Nombre d'occurences": month_counts.reindex(range(1, 13), fill_value=0).tolist()}

        filtered_data1 = combined_data[(combined_data['ÉLIGIBILITÉ_AU_DON.']=="Définitivement non-eligible")]
        filtered_data2 = combined_data[(combined_data['ÉLIGIBILITÉ_AU_DON.']=="Temporairement Non-eligible")]

        row1_cos = st.columns(2)
        create_metric_card(row1_cos[0],title="Number of Donors By Months",plot_function=lambda: plot_radar_chart(data, categories),width="100%") 
        create_metric_card(row1_cos[1],title="Number of Donors By Months and by eligibility status",plot_function=lambda: by_months_status(filtered_data),width="100%") 
        
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}</div> """,  unsafe_allow_html=True)
        row1_cos = st.columns(1)
        create_metric_card(row1_cos[0],title=get_text("Evolution of Number of Donors"),plot_function=lambda: number_line(filtered_data),width="100%") 
        
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}</div> """,  unsafe_allow_html=True)
        row1_cos = st.columns(1)
        create_metric_card(row1_cos[0],title=get_text("Evolution of Donations"),plot_function=lambda: heatmap(filtered_data),width="100%") 
        
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}</div> """,  unsafe_allow_html=True)
        row1_cos = st.columns(1)
        create_metric_card(row1_cos[0],title=get_text("Evolution of Donation"),plot_function=lambda: jour(load_data1(),height="400px"),width="100%") 

        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}</div> """,  unsafe_allow_html=True)

        row1_cos = st.columns(2)
        create_metric_card(row1_cos[0],title=get_text("Others  Non Availability purposes"),plot_function=lambda: generate_wordcloud(filtered_data['Autre raisons,  preciser'].dropna()),width="100%") 
        create_metric_card(row1_cos[1],title=get_text("Others Non Eligibility Purposes"),plot_function=lambda: generate_wordcloud(filtered_data[filtered_data['ÉLIGIBILITÉ_AU_DON.'] == 'Définitivement non-eligible']['Si autres raison préciser'].dropna()),width="100%") 
        
        row1_cos = st.columns(2)
        create_metric_card(row1_cos[1],title=get_text("Non Eligibility purposes"),plot_function=lambda:generate_wordcloud_and_barchart(filtered_data1),width="100%") 
        create_metric_card(row1_cos[0],title=get_text("Non Availability purposes"),plot_function=lambda:generate_wordcloud_and_barchart(filtered_data2),width="100%") 
         


    elif selected_item =="Eligibility and Profile" : 
        st.markdown("""
    <style>
        .card-title {
            font-weight: bold;
            font-size: 1.2em;
            text-align: center;
            padding: 10px 0;
            color: #8a2be2;
        }
        
        /* Make plots use full width of their containers */
        .js-plotly-plot, .plotly, .plot-container {
            width: 100% !important;
        }
        
        /* Remove default padding from columns to maximize space */
        div[data-testid="column"] {
            padding: 0 !important;
            margin: 0 !important;
        }
        
        /* Style for the streamlit columns inside the container */
        div[data-testid="stHorizontalBlock"] {
            width: 100% !important;
            gap: 5px;
        }
    </style>
    """, unsafe_allow_html=True)
        st.markdown(
        """
        <style>
            div[data-baseweb="select"] > div {
                min-height: 10px !important;  /* Adjust height */
            }
        </style>
        """,
        unsafe_allow_html=True
    )
        combined_data = get_combined_data()
        row1_cols = st.columns(3)
        with row1_cols[1]:
            all_arrondissement = combined_data['Arrondissement_de_résidence_'].unique()
            selected_arrondissement = st.sidebar.multiselect(
                get_text("Districts"),
                all_arrondissement
            )
            if not selected_arrondissement:
                selected_arrondissement = all_arrondissement
        with row1_cols[2]:
            all_case = combined_data['ÉLIGIBILITÉ_AU_DON.'].unique()
            selected_case = st.sidebar.multiselect(
                "Eligible",
                all_case # Limiter par défaut pour améliorer les performances
            )
            if not selected_case:
                selected_case = all_case
        # Filtrer les données avec tous les filtres
        filtered_data = combined_data[
            (combined_data['Arrondissement_de_résidence_'].isin(selected_arrondissement)) &
            (combined_data['ÉLIGIBILITÉ_AU_DON.'].isin(selected_case))]
        st.markdown("""
            <style>
            .block-container {
                padding-top: 0rem;
                padding-bottom: 0rem;
                padding-left: 1rem;
                padding-right: 1rem;
                margin-top: 0px;
            }
                
                /* Remove extra padding around header */
                header {
                    margin-bottom: 0rem !important;
                    padding-bottom: 0rem !important;
                }
            </style>
            """, unsafe_allow_html=True)
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("Welcome to the  Eligibility and Profile Section")}
        </div> """,  unsafe_allow_html=True)
        row1_cos = st.columns(1)
        create_metric_card(row1_cos[0],title=get_text("Eligibility Profile"),plot_function=lambda: circle(filtered_data),width="100%") 
        
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}</div> """,  unsafe_allow_html=True)
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}<br/></div> """,  unsafe_allow_html=True)
        if st.session_state.language == "English"  : 
            create_metric_card(row1_cos[0],title=get_text("Ideal Donor"),plot_function=lambda: display_ideal_chart_e() , width="100%") 
        else :  
            create_metric_card(row1_cos[0],title="Ideal Donor",plot_function=lambda: display_ideal_chart_f() , width="100%") 
        
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}</div> """,  unsafe_allow_html=True)
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("")}</div> """,  unsafe_allow_html=True)
        create_metric_card(row1_cos[0],title=get_text("Eligibility In Douala"),plot_function=lambda: three(filtered_data),width="100%") 


    elif selected_item == "Dataset Insights":
        st.markdown("""
    <style>
        .card-title {
            font-weight: bold;
            font-size: 1.2em;
            text-align: center;
            padding: 10px 0;
            color: #8a2be2;
        }
        
        /* Make plots use full width of their containers */
        .js-plotly-plot, .plotly, .plot-container {
            width: 100% !important;
        }
        
        /* Remove default padding from columns to maximize space */
        div[data-testid="column"] {
            padding: 0 !important;
            margin: 0 !important;
        }
        
        /* Style for the streamlit columns inside the container */
        div[data-testid="stHorizontalBlock"] {
            width: 100% !important;
            gap: 5px;
        }
    </style>
    """, unsafe_allow_html=True)
        st.markdown(
        """
        <style>
            div[data-baseweb="select"] > div {
                min-height: 10px !important;  /* Adjust height */
            }
        </style>
        """,
        unsafe_allow_html=True
    )
        st.markdown("""
            <style>
            .block-container {
                padding-top: 0rem;
                padding-bottom: 0rem;
                padding-left: 1rem;
                padding-right: 1rem;
                margin-top: 0px;
            }
                
                /* Remove extra padding around header */
                header {
                    margin-bottom: 0rem !important;
                    padding-bottom: 0rem !important;
                }
            </style>
            """, unsafe_allow_html=True)
        combined_data = get_combined_data()
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("Welcome to the  social demographics insights")}
        </div> """,  unsafe_allow_html=True)
        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;"> </div> """,  unsafe_allow_html=True)
        all_arrondissement = combined_data['Arrondissement_de_résidence_'].unique()
        selected_arrondissement = st.sidebar.multiselect(
            get_text("Districts"),
            all_arrondissement
        )
        if not selected_arrondissement:
            selected_arrondissement = all_arrondissement

        all_case = combined_data['ÉLIGIBILITÉ_AU_DON.'].unique()
        selected_case = st.sidebar.multiselect(
            "Eligible",
            all_case # Limiter par défaut pour améliorer les performances
        )
        if not selected_case:
            selected_case = all_case
        # Filtrer les données avec tous les filtres
        filtered_data = combined_data[
            (combined_data['Arrondissement_de_résidence_'].isin(selected_arrondissement)) &
            (combined_data['ÉLIGIBILITÉ_AU_DON.'].isin(selected_case))]

        cols = st.columns(2)
        create_metric_card(cols[0], title=get_text("Marital Status"), plot_function=lambda: render_frequency_pieh(filtered_data["Situation_Matrimoniale_(SM)"], legend_top="70%", legend_left="15%"), width="100%")
        create_metric_card(cols[1], title=get_text("Level"), plot_function=lambda: render_frequency_pieh(filtered_data["Niveau_d'etude"]), width="100%")
        cols = st.columns(2)
        create_metric_card(cols[0], title=get_text("Sector"), plot_function=lambda: render_frequency_pie(filtered_data["Secteur"], legend_left="2%", legend_top="2%"), width="100%")
        create_metric_card(cols[1], title=get_text("Gender"), plot_function=lambda: render_frequency_pieh(filtered_data["Genre_"], legend_top="80%", legend_left="65%"), width="100%")
        
        cols = st.columns(2)
        create_metric_card(cols[0], title=get_text("Profession"), plot_function=lambda: render_frequency_pie(filtered_data["categories"], legend_left="2%", legend_top="2%"), width="100%")
        create_metric_card(cols[1], title=get_text("Eligibility Status"), plot_function=lambda: render_frequency_pieh(filtered_data["ÉLIGIBILITÉ_AU_DON."], legend_top="90%", legend_left="15%"), width="100%")

        cols = st.columns(2)
        option = compute_age_distribution(filtered_data['Age'])
        create_metric_card(cols[0], title=get_text('Age distribution'), plot_function=lambda: st_echarts(options=option, height=400, width="100%"), width="100%")
        create_metric_card(cols[1], title=get_text('Population Pyramid'), plot_function=lambda: plot_age_pyramid(filtered_data, height=400), width="100%")

    elif selected_item == "Options":
        st.markdown(
            """
            <h1 style='color: red; text-align: center; margin-top: 0;'>Options</h1>
            """,
            unsafe_allow_html=True
        )
        st.markdown(get_text('### Theme Personalization'))
        # Initialize theme in session state
        if 'theme' not in st.session_state:
            st.session_state.theme = "Light"

        # Define themes with fonts and colors
        THEMES = {
            "Light": {
                "primaryColor": "#F50307",
                "backgroundColor": "#FBFBFB",
                "secondaryBackgroundColor": "#EC8282",
                "textColor": "#000000",
                "font": "sans serif"
            },
            "Dark": {
                "primaryColor": "#f50307",
                "backgroundColor": "#0E1117",
                "secondaryBackgroundColor": "#2D2A2A",
                "textColor": "#fff4f4",
                "font": "sans serif"
            },
            "Blue": {
                "primaryColor": "#087CEF",
                "backgroundColor": "#FFFFFF",
                "secondaryBackgroundColor": "#7FD6AC",
                "textColor": "black",
                "font": "sans serif"
            }
        }   
        @st.cache_resource 
        def apply_theme(theme_choice):
            # Get the selected theme
            selected_theme = THEMES[theme_choice]

            # Update session state
            st.session_state.theme = theme_choice

            # Ensure the .streamlit directory exists
            config_dir = os.path.join(os.path.expanduser("~"), ".streamlit")
            os.makedirs(config_dir, exist_ok=True)

            # Write the selected theme to the config.toml file
            config_path = os.path.join(config_dir, "config.toml")
            with open(config_path, "w") as config_file:
                config_file.write(f"""
                [theme]
                primaryColor = "{selected_theme['primaryColor']}"
                backgroundColor = "{selected_theme['backgroundColor']}"
                secondaryBackgroundColor = "{selected_theme['secondaryBackgroundColor']}"
                textColor = "{selected_theme['textColor']}"
                font = "{selected_theme['font']}"
                """)
            # Show success message
            st.success(f"Theme will be changed to {theme_choice}! Click Again to confirm. ")
    

        # Theme selection dropdown
        theme = st.selectbox(
            get_text("Choose your theme"), 
            list(THEMES.keys()),
            index=list(THEMES.keys()).index(st.session_state.theme),
            key="widget_theme"
        )

        # Get selected theme colors
        selected_theme = THEMES[st.session_state.theme]

        # Apply button with dynamic text color
        apply_button = st.markdown(
            f"""
            <style>
                div.stButton > button {{
                    color: {selected_theme['textColor']} !important;
                    border-radius: 5px;
                    padding: 10px 20px;
                    font-size: 16px;
                }}
            </style>
            """,
            unsafe_allow_html=True
        )

        # Add confirmation button to apply theme
        if st.button(get_text("Apply Theme"), key="apply_theme_button"):
            apply_theme(st.session_state.widget_theme)
    #__________________________________
        # Initialize language if not set
        if 'language' not in st.session_state:
            st.session_state.language = list(TRANSLATIONS.keys())[0]  # Default to first language

        # Function to apply language
    
        def apply_language():
            st.session_state.language = st.session_state.widget_language

            
        # Language selection with callback
        st.selectbox(
            get_text("Choose your language"), 
            list(TRANSLATIONS.keys()),
            index=list(TRANSLATIONS.keys()).index(st.session_state.language),
            key="widget_language"
        )

        if st.button(get_text("Apply Language"), key="apply_language_button"):
            st.success(f"Language will be changed to {st.session_state.widget_language  }! Click Again to confirm. ")
            apply_language()
        col=st.columns(2)
        with col[1]:
            st.image("OIP.jpeg", width=500)
                    
    elif selected_item == "Home":
        site_url = "https://githubsssssssssssssssssssssss-challenge-es-zdqvvx.streamlit.app/"
        qr = qrcode.QRCode(version=1, box_size=10, border=4)
        qr.add_data(site_url)
        qr.make(fit=True)
        img = qr.make_image(fill_color="black", back_color="white")
        buffered = io.BytesIO()
        img.save(buffered, format="PNG")
        st.sidebar.image(buffered)
    
        col = st.columns(1)

        with col[0]: 
            st.markdown(
        f"""
        <div class="metric-box" style="background-color: #FFFBF5; height : 80px">
            <div class="metric-title" style="text-align: center; font-size: 29px; font-weight: bold; color: red;">{get_text("Welcome to Our Blood Donation Campaign Dashboard")}
                    </div>
</div>
           
        """,
        unsafe_allow_html=True
    )           

            
        st.markdown("""
            <style>
            .block-container {
                padding-top: 2rem;
                padding-bottom: 0rem;
                padding-left: 5rem;
                padding-right: 2rem;
                margin-top: 0px;
            }
                
                /* Remove extra padding around header */
                header {
                    margin-bottom: 0rem !important;
                    padding-bottom: 0rem !important;
                }
            </style>
            """, unsafe_allow_html=True)
        st.markdown(
    """
    <style>
    .sidebar .sidebar-content {
        background-color: #2C3E50;
        color: white;
    }
    .metric-box {
        border-radius: 30px;
        padding: 1px;
        margin: 10px;
        text-align: center;
        box-shadow: 0 15px 40px rgba(0,0,0,0.3);
        transition: all 0.5s ease;
        cursor: pointer;
        position: relative;
        overflow: hidden;
        height: 160px;
        display: flex;
        flex-direction: column;
        justify-content: center;
        border: 2px solid rgba(255, 255, 255, 0.3);
    }
    .metric-box:hover {
        transform: scale(1.1) rotate(0deg);
        box-shadow: 0 20px 50px rgba(0,0,0,0.4);
    }
    .metric-box h1 {
        font-size: 40px;
        color: white;
        text-shadow: 2px 2px 6px rgba(0,0,0,0.4);
    }
    .metric-box p {
        font-size: 22px;
        color: white;
        margin: 15px 0 0 0;
        text-shadow: 1px 1px 4px rgba(0,0,0,0.3);
        z-index: 1;
        font-weight: bold;
    }
        </style>
    """,
    unsafe_allow_html=True
)

        row1_cols = st.columns([1,1,1,1])
        with row1_cols[0]:
            
            Volonteers = load_data1().shape[0] +load_data2(get_modification_time()).shape[0]
            
            st.markdown(f"""
            <div class="metric-box " style="background: linear-gradient(135deg, #6D8299, #A3BFFA)";>
                <h1>{Volonteers}💪</h1>
                <p><strong >{get_text("Volunteers")}</strong></p>
            </div>
            """, unsafe_allow_html=True)
    
        with row1_cols[1]:
            df1 = load_data1()
            df2 = load_data2(get_modification_time())
            eligible = df1[df1["ÉLIGIBILITÉ_AU_DON."] == "Eligible"].shape[0] + df2[df2["ÉLIGIBILITÉ_AU_DON."] == "Eligible"].shape[0]
            st.markdown(f"""
        <div class="metric-box " style="background: linear-gradient(135deg, #2E7D32, #81C784)";>
            <h1>{eligible} ✅</h1>
            <p><strong>{get_text("Eligibles")}</strong></p>
        </div>
        """,
        unsafe_allow_html=True
    )

        with row1_cols[2]:
            T_eligible = df1[df1["ÉLIGIBILITÉ_AU_DON."] == "Temporairement Non-eligible"].shape[0] + df2[df2["ÉLIGIBILITÉ_AU_DON."] == "Temporairement Non-eligible"].shape[0]
            st.markdown(f"""
        <div class="metric-box" style="background: linear-gradient(135deg, #F9A825, #FFCA28)">
            <h1>{T_eligible} ⚠️</h1>
            <p><strong>{get_text("Temporarily Not Eligible")}</strong></p>
        </div>
        """,
        unsafe_allow_html=True
    )

        with row1_cols[3]:
            T_eligible = df1[df1["ÉLIGIBILITÉ_AU_DON."] == "Définitivement non-eligible"].shape[0] + df2[df2["ÉLIGIBILITÉ_AU_DON."] == "Définitivement non-eligible"].shape[0]
            st.markdown(
        f"""
        <div class="metric-box " style="background: linear-gradient(135deg, #C2185B, #F06292)">
            <h1>{T_eligible} ❌</h1>
            <p><strong>{get_text("Definitely Non-eligible")}</strong></p>
        </div>
        """,
        unsafe_allow_html=True
    )

        row1_cols = st.columns([0.8,0.7,1])
        with row1_cols[1]:
            df =  get_combined_data()
            nomb = df['Si oui preciser la date du dernier don'].count()
            st.markdown(
        f"""
         <div class="metric-box" style="background: linear-gradient(135deg, #ED9ED6, #C683D7)">
            <h1>{ nomb}❤️</h1>
            <p><strong>{get_text("Have Ever Donated")}</strong></p>
        </div>
        """,
        unsafe_allow_html=True
    )

        with row1_cols[0]:
            df= load_data1()[(load_data1()['ÉLIGIBILITÉ_AU_DON.']=="Définitivement non-eligible")]
            df = df.dropna(subset=["Raison_indisponibilité_fusionnée"])
            df['Raison_indisponibilité_fusionnée'] = df["Raison_indisponibilité_fusionnée"].astype(str)
            df['Raison_indisponibilité_fusionnée'] = df["Raison_indisponibilité_fusionnée"].str.split(';')
            df_exploded = df.explode("Raison_indisponibilité_fusionnée")
            df_exploded = df_exploded.dropna(subset=["Raison_indisponibilité_fusionnée"])
            element_frequent = df_exploded["Raison_indisponibilité_fusionnée"].mode()[0]
            st.markdown(
        f"""
        <div class="metric-box" style="background: linear-gradient(135deg, #FFC436, #D2DE32)">
            <h1>{element_frequent} 🩺</h1>
            <p><strong>{get_text("Most Non eligibility Purpose")}</strong></p>
        </div>
        """,
        unsafe_allow_html=True
    )

        with row1_cols[2]:
            df= load_data1()[(load_data1()['ÉLIGIBILITÉ_AU_DON.']=="Temporairement Non-eligible")]
            df = df.dropna(subset=["Raison_indisponibilité_fusionnée"])
            df['Raison_indisponibilité_fusionnée'] = df["Raison_indisponibilité_fusionnée"].astype(str)
            df['Raison_indisponibilité_fusionnée'] = df["Raison_indisponibilité_fusionnée"].str.split(';')
            df_exploded = df.explode("Raison_indisponibilité_fusionnée")
            df_exploded = df_exploded.dropna(subset=["Raison_indisponibilité_fusionnée"])
            element_frequent = df_exploded["Raison_indisponibilité_fusionnée"].mode()[0]
            st.markdown(
        f"""
        <div class="metric-box " style="background: linear-gradient(15deg, #AED2FF, #6F42C1)">
            <h1 style="font-size: 33px">{element_frequent} 🩸 </h1>
            <p><strong>{get_text("Most Non Availability Purpose")}</strong></p>
        </div>
        """,
        unsafe_allow_html=True
    )

        gdf = load_shapefile("gadm41_CMR_0.shp")
        data_df_3 = get_preprocessed_data(3)
        data_df_2 = get_preprocessed_data(2)

        row=st.columns(1)
        with row[0]:
            with stylable_container(
                    key='944',
                    css_styles=f"""
                        {{
                        width : 100%;
                        border: 1px solid #c0c0c0;
                        border-radius: 10px;
                        flex-direction: column;
                        background-color: #f8f9fa;
                        box-shadow: 0px 4px 6px  2px rgba(0, 0, 0, 0.8);
                        }}
                        
                        .card-title {{
                                font-weight: bold;
                                margin: 0px;
                                padding: 0px;
                                font-size: 1em;
                                text-align: center;
                                color: #8a2be2;  # Light purple color
                            }}
                    """ ):
                    st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 22px; font-weight: bold; color:black">{get_text("Keys Rate")}</div> """,  unsafe_allow_html=True)       
                    row = st.columns([1,1,1])
                    with row[0]:
                        option_jauge = {
                        "tooltip": {
                            "formatter": '{a} <br/>{b} : {c}%'
                        },
                        "series": [
                            {
                                "name": "Progression",
                                "type": "gauge",
                                "startAngle": 180,
                                "endAngle": 0,
                                "radius": "90%",
                                "itemStyle": {
                                    "color": "red",
                                    "shadowColor": "rgba(0,138,255,0.45)",
                                    "shadowBlur": 10,
                                    "shadowOffsetX": 2,
                                    "shadowOffsetY": 2
                                },
                                "progress": {
                                    "show": True,
                                    "roundCap": True,
                                    "width": 10
                                },
                                "pointer": {
                                    "length": "60%",
                                    "width": 3,
                                    "offsetCenter": [0, "5%"]
                                },
                                "detail": {
                                    "valueAnimation": True,
                                    "formatter": "{value}%",
                                    "backgroundColor": "red",
                                    "width": "100%",
                                    "lineHeight": 25,
                                    "height": 16,
                                    "borderRadius": 188,
                                    "offsetCenter": [0, "40%"],
                                    "fontSize": 18
                                },
                                "data": [{
                                    "value": round(100*eligible/Volonteers, 1),  # Example value
                                    "name": get_text("Eligibility Rate"),  # Label for the value"Eligibility Rate"
                                }]
                            }
                        ]
                    }
                        st_echarts(options=option_jauge, key="1")
                    with row[1]:
                        option_jauge = {
                        "tooltip": {
                            "formatter": '{a} <br/>{b} : {c}%'
                        },
                        "series": [
                            {
                                "name": "Progression",
                                "type": "gauge",
                                "startAngle": 180,
                                "endAngle": 0,
                                "radius": "90%",
                                "itemStyle": {
                                    "color": "red",
                                    "shadowColor": "rgba(0,138,255,0.45)",
                                    "shadowBlur": 10,
                                    "shadowOffsetX": 2,
                                    "shadowOffsetY": 2
                                },
                                "progress": {
                                    "show": True,
                                    "roundCap": True,
                                    "width": 10
                                },
                                "pointer": {
                                    "length": "60%",
                                    "width": 3,
                                    "offsetCenter": [0, "5%"]
                                },
                                "detail": {
                                    "valueAnimation": True,
                                    "formatter": "{value}%",
                                    "backgroundColor": "red",
                                    "width": "100%",
                                    "lineHeight": 25,
                                    "height": 16,
                                    "borderRadius": 188,
                                    "offsetCenter": [0, "40%"],
                                    "fontSize": 18
                                },
                                "data": [{
                                    "value": round(100*T_eligible/Volonteers, 1),  # Example value
                                    "name": get_text("Temporarily Eligibility Rate"), #"Temporarily Eligibility Rate"
                                }]
                            }
                        ]
                    }
                        st_echarts(options=option_jauge, key="2")
                    with row[2]:
                        option_jauge = {
                        "tooltip": {
                            "formatter": '{a} <br/>{b} : {c}%'
                        },
                        "series": [
                            {
                                "name": "Progression",
                                "type": "gauge",
                                "startAngle": 180,
                                "endAngle": 0,
                                "radius": "90%",
                                "itemStyle": {
                                    "color": "red",
                                    "shadowColor": "rgba(0,138,255,0.45)",
                                    "shadowBlur": 10,
                                    "shadowOffsetX": 2,
                                    "shadowOffsetY": 2
                                },
                                "progress": {
                                    "show": True,
                                    "roundCap": True,
                                    "width": 10
                                },
                                "pointer": {
                                    "length": "60%",
                                    "width": 3,
                                    "offsetCenter": [0, "5%"]
                                },
                                "detail": {
                                    "valueAnimation": True,
                                    "formatter": "{value}%",
                                    "backgroundColor": "red",
                                    "width": "100%",
                                    "lineHeight": 25,
                                    "height": 16,
                                    "borderRadius": 188,
                                    "offsetCenter": [0, "40%"],
                                    "fontSize": 18
                                },
                                "data": [{
                                    "value": round(100*(Volonteers-eligible-T_eligible)/Volonteers, 1),  # Example value
                                    "name": get_text("Non Eligibility Rate"), #"Ilegibility Rate"
                                }]
                            }
                        ]
                    }
                        st_echarts(options=option_jauge, key="3")
                    

        row = st.columns([1.1,1.2])
        with row[0] : 
            st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 20px; font-weight: bold; color:black">{get_text("Donnor Insights")}</div> """,  unsafe_allow_html=True)       
            with stylable_container(
                key='1116',
                css_styles=f"""
                    {{
                    width : 100%;
                    border: 1px solid #c0c0c0;
                    border-radius: 10px;
                    flex-direction: column;
                    background-color: #f8f9fa;
                    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.8);
                    }}
                    
                    .card-title {{
                            font-weight: bold;
                            margin: 0px;
                            padding: 0px;
                            font-size: 1em;
                            text-align: center;
                            color: #8a2be2;  # Light purple color
                        }}
                """ ):
                    row_1 = st.columns(2)
                    st.write("")
                    with row_1[0] :
                        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 20px; font-weight: bold; color:red">{get_text("Blood distribution")}</div> """,  unsafe_allow_html=True)       
                        pie_data = count_frequencies(load_data3()["Groupe Sanguin ABO / Rhesus"])
                        render_frequency_pie2(pie_data)
                    with row_1[1] :
                        st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 20px; font-weight: bold; color:red">{get_text("Donation Type")}</div> """,  unsafe_allow_html=True)       
                        pie_data = count_frequencies(load_data3()["Type de donation"])
                        render_frequency_pie2(pie_data)
                    st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 20px; font-weight: bold; color:red">{get_text("Donor Profile")}</div> """,  unsafe_allow_html=True)       
                    if st.session_state.language == "English":
                        display_ideal_chart_e() 
                    else : display_ideal_chart_f()
            st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 20px; font-weight: bold; color:black">{get_text("Social and Demographics stats")}</div> """,  unsafe_allow_html=True)       
            with stylable_container(
                key='1152',
                css_styles=f"""
                    {{width : 100%;
                    
        
                        border: 1px solid #c0c0c0;
                    border-radius: 10px;
                    margin: 0px;  # Small margin for spacing
                    padding: 0px;
                    display: flex;
                    flex-direction: column;
                    background-color: #f8f9fa;
                    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.811111);
                    transition: all 0.2s ease;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    color: green;
                    }}
                    
                    .card-title {{
                            font-weight: bold;
                            margin: 0px;
                            padding: 0px;
                            font-size: 1em;
                            text-align: center;
                            color: #8a2be2;  # Light purple color
                        }}
                """ ):
                row_1 = st.columns([1,1.4])
                with row_1[0] :
                    st.markdown(f'''<div style='text-align: center;'>{get_text("Level")}</div>''', unsafe_allow_html=True)
                    pie_data = count_frequencies(load_data1()["Niveau_d'etude"])
                    render_frequency_pie2(pie_data)
                    st.markdown(f'''<div style='text-align: center;'>{get_text("Marital Status")}</div>''', unsafe_allow_html=True)
                    pie_data = count_frequencies(load_data1()["Situation_Matrimoniale_(SM)"])
                    render_frequency_pie2(pie_data)
                with row_1[1] :
                    st.markdown(f'''<div style='text-align: center;'>{get_text("Gender")}</div>''', unsafe_allow_html=True)
                    pie_data = count_frequencies(load_data1()["Genre_"])
                    render_frequency_pie2(pie_data)
                    st.markdown(f'''<div style='text-align: center;'>{get_text("Eligibility")}</div>''', unsafe_allow_html=True)
                    pie_data = count_frequencies(load_data1()["ÉLIGIBILITÉ_AU_DON."])
                    render_frequency_pie2(pie_data)



        with row[1] : 
            st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 20px; font-weight: bold; color:black">{get_text("Cartography Summary")}</div> """,  unsafe_allow_html=True)       
            with stylable_container(
                key='1201',
                css_styles=f"""
                    {{width : 100%;
                        border: 1px solid #c0c0c0;
                    border-radius: 10px;
                    margin: 0px;  # Small margin for spacing
                    padding: 0px;
                    display: flex;
                    flex-direction: column;
                    background-color: #f8f9fa;
                    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.811111);
                    transition: all 0.2s ease;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    color: green;
                    }}
                    
                    .card-title {{
                            font-weight: bold;
                            margin: 0px;
                            padding: 0px;
                            font-size: 1em;
                            text-align: center;
                            color: #8a2be2;  # Light purple color
                        }}
                """ ):
                    st.write("")
                    row_1 = st.columns([1,0.8])
                    with row_1[1] :
                        st.dataframe(data_df_3.head(5), use_container_width=True)
                        st.dataframe(data_df_2.head(8), use_container_width=True)
                    with row_1[0] :
                        gdf = load_shapefile("gadm41_CMR_0.shp")
                        m = folium.Map(location=[10.87, 13.52], zoom_start=5, control_scale=True, tiles="CartoDB positron", height=901)
                        folium.GeoJson(
                            gdf.to_crs(epsg=4326).__geo_interface__,
                            name="Cameroun",
                            style_function=lambda x: {'fillColor': 'transparent', 'color': 'black', 'weight': 2, 'fillOpacity': 0.0}
                        ).add_to(m)
                        m.fit_bounds([[6.2, 9.5], [0, 16.75]])

                        combined_data = get_hierarchical_data()
                        geo_data_3 = load_shapefile("gadm41_CMR_3.shp", simplify_tolerance=0.01)

                        geo_data_3['centroid_lat'] = geo_data_3['geometry'].centroid.y
                        geo_data_3['centroid_lon'] = geo_data_3['geometry'].centroid.x

                        centroid_dict = dict(zip(geo_data_3['NAME_3'], zip(geo_data_3['centroid_lat'], geo_data_3['centroid_lon'])))

                        geo_data_3 = geo_data_3.merge(combined_data, left_on='NAME_3', right_on='Arrondissement_de_résidence_')
                        mask_missing = combined_data['latitude'].isna() | combined_data['longitude'].isna()
                        missing_indices = combined_data[mask_missing].index

                        for idx in missing_indices:
                            arrondissement = combined_data.loc[idx, 'Arrondissement_de_résidence_']
                            if arrondissement in centroid_dict:
                                combined_data.loc[idx, 'latitude'] = centroid_dict[arrondissement][0]
                                combined_data.loc[idx, 'longitude'] = centroid_dict[arrondissement][1]

                        combined_data = combined_data.dropna(subset=['latitude', 'longitude'])
                        quartier_counts = combined_data['Quartier_de_Résidence_'].value_counts().reset_index()
                        quartier_counts.columns = ['Quartier_de_Résidence_', 'count']

                        quartier_locations = combined_data.groupby('Quartier_de_Résidence_').agg({
                            'latitude': 'first',
                            'longitude': 'first'
                        }).reset_index()

                        quartier_data = quartier_counts.merge(quartier_locations, on='Quartier_de_Résidence_')

                        min_count = quartier_data['count'].min() if not quartier_data.empty else 1
                        max_count = quartier_data['count'].max() if not quartier_data.empty else 1
                        for _, row in quartier_data.iterrows():
                            quartier = row["Quartier_de_Résidence_"]
                            count = row["count"]

                            radius = 3 + 4 * ((count - min_count) / max(max_count - min_count, 1)) * 10

                            folium.CircleMarker(
                                location=[row["latitude"], row["longitude"]],
                                radius=radius,
                                tooltip=f"{quartier}: {count} {get_text('Volonteers')}",
                                color='red',
                                fill=True,
                                fill_color='red'
                            ).add_to(m)

                        folium.LayerControl().add_to(m)

                        folium_static(m, width=400, height=500)

            st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 20px; font-weight: bold; color:black">{get_text("Campaign efficacity")}</div> """,  unsafe_allow_html=True)       
            with stylable_container(
                key='1294',
                css_styles=f"""
                    {{width : 100%;
                    
                        border: 1px solid #c0c0c0;
                    border-radius: 10px;
                    margin: 0px;  # Small margin for spacing
                    padding: 0px;
                    display: flex;
                    flex-direction: column;
                    background-color: #f8f9fa;
                    box-shadow: 0 4px 6px rgba(0, 0, 0, 0.811111);
                    transition: all 0.2s ease;
                    overflow: hidden;
                    text-overflow: ellipsis;
                    color: green;
                    }}
                    
                    .card-title {{
                            font-weight: bold;
                            margin: 0px;
                            padding: 0px;
                            font-size: 1em;
                            text-align: center;
                            color: #8a2be2;  # Light purple color
                        }}
                """ ):
                    st.write("")
                    df = load_data1()
                    df['Date de remplissage de la fiche'] = pd.to_datetime(df['Date de remplissage de la fiche'])
                    m= df['Date de remplissage de la fiche'].dt.month
                    month_counts = m.value_counts().sort_index()
                    categories = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
                    data = {"Nombre d'occurences": month_counts.reindex(range(1, 13), fill_value=0).tolist()}
                    jour(load_data1(),height="400px")
                    number_line(load_data1())  

        row1_cos = st.columns(1)
        create_metric_card(row1_cos[0],title=get_text("Non eligibility and Unavailability purpose distribution"),plot_function=lambda: women_reasons(df = get_combined_data()),height=500) 

        
    elif selected_item == "Donations":
        with st.sidebar:
            donations_button = st.sidebar.button("Volunteer registration", use_container_width=True)
            dataset_button = st.sidebar.button("Dataset", use_container_width=True)
            #donations_button = st.sidebar.button("Donor registration", use_container_width=True)
            
    # Donor Registration page
        if not dataset_button :
            geolocator = Nominatim(user_agent="geoapi")
            st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
        {get_text("REGISTRATION")}
        </div> """,  unsafe_allow_html=True)

            with st.form("donor_registration_form"):
                col1, col2 = st.columns(2)
                last_donor_number = 0
                # Liste des pays d'Afrique centrale
                pays_afrique_centrale = ["Cameroun", "République centrafricaine", "Tchad", "République du Congo", "République démocratique du Congo", "Gabon", "Guinée équatoriale", "São Tomé-et-Príncipe", "Autres"]
                if not st.session_state.donors.empty:
                    last_donor_number = int(st.session_state.donors['id'].str.extract(r'(\d+)').max())
                elif not load_data2(get_modification_time()).empty:
                    last_donor_number = int(load_data2(get_modification_time())['ID'].str.extract(r'(\d+)').max())
                else:
                    last_donor_number = 0
                with col1:
                    import datetime
                    donor_number = st.number_input(get_text("Registration Number"), min_value=last_donor_number + 1, value=last_donor_number + 1, step=1)
                    ID = f"DONOR_{donor_number}"
                    Date_remplissage = st.date_input(get_text("Filling Date"))
                    Date_naiss = st.date_input(get_text("BirthDate"), min_value=datetime.date(1960, 1, 1))
                    st.write(f"Âge : {(Date_remplissage - Date_naiss).days //365} ans")
                    Niveau_detude = st.selectbox(get_text("Education Level"), ["Primary", "Secondary", "High School", "Bachelor's", "Master's", "PhD", "Other"])
                    Genre_ = st.radio(get_text("Gender"), ["Homme", "Femme"])
                    Taille_ = st.number_input(get_text("Height"), min_value=100, max_value=220, step=1)
                    Poids = st.number_input(get_text("Weight"), min_value=60, max_value=200, step=1)
                    Profession_ = st.text_input("Profession")  
                    
                                    
                with col2:
                    A_deja_donne_le_sang_ = st.radio(get_text("Has already donated blood"), ["Yes", "No"]) 
                    geo_data_3 = gpd.read_file("gadm41_CMR_3.shp")
                    districts = geo_data_3["NAME_3"].unique()
                    Arrondissement_de_residence_ = st.selectbox(
                        get_text ("District of Residence"),
                        options=districts,
                        index=0,  # Sélectionner le premier élément par défaut
                        help="Sélectionnez votre district de résidence"
                    )
                            
                    Quartier_de_Residence_ = st.text_input( get_text("residential neighborhood"))

                    # Search button for location
                    if Quartier_de_Residence_:
                        location = geolocator.geocode(f"{Quartier_de_Residence_}, Cameroun")
                        if location:
                            latitude = location.latitude 
                            longitude = location.longitude
                            Quartier_de_Residence_ = geolocator.geocode(f"{Quartier_de_Residence_}, Cameroun")
                            st.success(f"📍 {Quartier_de_Residence_} {get_text('Localised')} → Latitude: {round(location.latitude,2)}, Longitude: {round(location.longitude,2)}")
                        else:
                            st.error(f"❌ {get_text('Location not found in Cameroon.')}")
                    if st.form_submit_button(get_text("verify")):
                        pass
                    Nationalite_ = st.selectbox(get_text("Country"), options=pays_afrique_centrale)
                    Age =(Date_remplissage - Date_naiss).days //365
                    Religion_ = st.selectbox(get_text("Religion"), ["Christianity", "Islam", "Judaism", "Buddhism", "Other", "No religion"])
                    Situation_Matrimoniale_SM = st.selectbox(get_text("Marital Status"), ["Single", "Married", "Divorced", "Widowed", "Domestic Partnership"])
                    if A_deja_donne_le_sang_ == "Yes":
                        Date_dernier_don_ = st.date_input(get_text("If already donated, specify the date of the last donation"))
                    else:
                        Date_dernier_don_ = None  
                    Taux_dhemoglobine_ = st.number_input(get_text("Hemoglobin Level"), min_value=0.0, max_value=25.0, step=0.1)               
                    #ELIGIBILITE_AU_DON = st.selectbox(get_text("ELIGIBILITY FOR DONATION"), ["Not Eligible", "Eligible"])
                # Create a container for non-eligibility fields that will be shown/hidden based on eligibility
                non_eligibility_container = st.container()
                
                # Use an empty element to control visibility, which will be filled if Not Eligible is selected
                with non_eligibility_container:
                    col3, col4 ,col5= st.columns(3)
                    with col3:
                            Est_sous_anti_biotherapie = st.radio(get_text("Under antibiotic therapy?"), ["No", "Yes"])
                            Taux_dhemoglobine_bas = st.radio(get_text("Low hemoglobin level?"), ["No", "Yes"])
                            Date_dernier_don_3_mois = st.radio(get_text("Last donation date < 3 months?"), ["No", "Yes"])
                            IST_recente = st.radio(get_text("Recent STI (Excluding HIV, Hbs, Hcv)?"), ["No", "Yes"])
                            DDR_incorrecte = st.radio(get_text("Incorrect DDR if < 14 days before donation?"), ["No", "Yes"])
                            Allaitement = st.radio(get_text("Breastfeeding?"), ["No", "Yes"])
                    with col4:    
                            Accouchement_6mois = st.radio(get_text("Gave birth in the last 6 months?"), ["No", "Yes"])
                            Interruption_grossesse = st.radio(get_text("Pregnancy termination in the last 6 months?"), ["No", "Yes"])
                            Enceinte = st.radio(get_text("Pregnant?"), ["No", "Yes"])
                            Antecedent_transfusion = st.radio(get_text("History of transfusion?"), ["No", "Yes"])
                            Porteur_HIV_HBS_HCV = st.radio(get_text("Carrier (HIV, Hbs, Hcv)?"), ["No", "Yes"])
                            Opere = st.radio(get_text("Operated?"), ["No", "Yes"])
                    with col5:
                            Drepanocytaire = st.radio(get_text("Sickle cell?"), ["No","Yes", ])
                            Diabetique = st.radio(get_text("Diabetic?"), [ "No","Yes"])
                            Hypertendus = st.radio(get_text("Hypertensive?"), ["No", "Yes"])
                            Asthmatiques = st.radio(get_text("Asthmatic?"), ["No", "Yes"])
                            Cardiaque = st.radio(get_text("Cardiac?"), ["No", "Yes"])
                            Tatoue = st.radio(get_text("Tattooed?"), ["No", "Yes"])
                            Scarifie = st.radio(get_text("Scarified?"), ["No", "Yes"])

                    Numero =  st.text_input(get_text("Numéro de téléphone"))  
                    Si_autres_raison = st.text_input(get_text("If other reasons, specify"))
                        
                if st.form_submit_button(get_text("Submit")):
                    if not Age or not Arrondissement_de_residence_ :
                        st.error("Please fill all required fields. It seems that you don't fill the date of birth")
                    else:
                        new_donor = pd.DataFrame({
                        'id': [f"DONOR_{donor_number}"],
                        'age': [Age],
                        'Date_remplissage': [Date_remplissage.strftime('%Y-%m-%d')],
                        'Date_naiss': [Date_naiss.strftime('%Y-%m-%d')],
                        'niveau_detude': [Niveau_detude],
                        'genre': [Genre_],
                        'taille': [Taille_],
                        'poids': [Poids],
                        'situation_matrimoniale': [Situation_Matrimoniale_SM],
                        'profession': [Profession_],
                        'arrondissement_residence': [Arrondissement_de_residence_],
                        'quartier_residence': [Quartier_de_Residence_],
                        'nationalite': [Nationalite_],
                        'religion': [Religion_],
                        'deja_donne_sang': [A_deja_donne_le_sang_],
                        'date_dernier_don': [Date_dernier_don_ if A_deja_donne_le_sang_ == 'Yes' else ''],
                        'taux_dhemoglobine': [Taux_dhemoglobine_],
                        'eligibilite_au_don': [None],
                        'lattitude': [latitude],
                        'longitude': [longitude],
                        'est_sous_anti_biotherapie': [Est_sous_anti_biotherapie ],
                        'taux_dhemoglobine_bas': [Taux_dhemoglobine_bas ],
                        'date_dernier_don_3_mois': [Date_dernier_don_3_mois ],
                        'ist_recente': [IST_recente ],
                        'ddr_incorrecte': [DDR_incorrecte ],
                        'allaitement': [Allaitement ],
                        'accouchement_6mois': [Accouchement_6mois ],
                        'interruption_grossesse': [Interruption_grossesse ],
                        'enceinte': [Enceinte ],
                        'antecedent_transfusion': [Antecedent_transfusion ],
                        'porteur_hiv_hbs_hcv': [Porteur_HIV_HBS_HCV ],
                        'opere': [Opere ],
                        'drepanocytaire': [Drepanocytaire ],
                        'diabetique': [Diabetique ],
                        'hypertendus': [Hypertendus ],
                        'asthmatiques': [Asthmatiques ],
                        'cardiaque': [Cardiaque ],
                        'tatoue': [Tatoue ],
                        'scarifie': [Scarifie ],
                        'autres_raisons': [Si_autres_raison ]
                    })
                        new_donor['eligibilite_au_don'] = check_eligibility(new_donor)['eligibility']
                        new_donor['raison'] = [check_eligibility(new_donor)['reasons']]
                        st.session_state.donors = pd.concat([st.session_state.donors, new_donor], ignore_index=True)
                        st.success(f"Thank you! Donor {ID} has been successfully registered.")
    #________________________________________
                        excel_file = "donnees.xlsx"

                        # Charger le fichier Excel existant
                        book = load_workbook(excel_file)

                        # Vérifier si 'Feuil1' existe
                        if 'Feuil1' in book.sheetnames:
                            with pd.ExcelWriter(excel_file, engine='openpyxl', mode='a', if_sheet_exists='overlay') as writer:
                                # Écrire le DataFrame à partir de la première ligne vide
                                startrow = book['Feuil1'].max_row  # Trouver la première ligne vide
                                new_donor.to_excel(writer, index=False, header=False, startrow=startrow, sheet_name='Feuil1')
        #_________________________________________
                    
                    # Check eligibility and display message 
                        if check_eligibility(new_donor)['eligibility'] == 'Eligible':
                            if st.session_state.language == "English" :
                                message("🎉🎆🎊 The Volonteer is  Eligible 🎉🎆🎊")
                                time.sleep(2)  
                                st.balloons()
                            else : 
                                message("🎉🎆🎊Le volontaire est  Eligible 🎉🎆🎊")
                                time.sleep(2)  
                                st.balloons()

                        elif  check_eligibility(new_donor)['eligibility'] =='Temporairement Non-eligible':
                            if st.session_state.language == "English" :
                                message(f"""The Volonteer is  Temporarily Non Eligible  😔😔""")
                            else : 
                                message(f"""Le volontaire est  Temporairement Non-eligible😔😔""")
                            st.markdown(f"<div style='text-align: center;'>- Reasons : </div>", unsafe_allow_html=True)
                            for reason in check_eligibility(new_donor)['reasons']:
                                st.markdown(f"<div style='text-align: center;'>- {reason}</div>", unsafe_allow_html=True)
                        else : 
                            if st.session_state.language == "English" :
                                message(f"""The Volonteer is Definitely Non Eligible  😔😔""")
                            else : 
                                message(f"""Le volontaire est  Definitivement Non-eligible😔😔""")
                            st.markdown(f"<div style='text-align: center;'>- Reasons : </div>", unsafe_allow_html=True)
                            for reason in check_eligibility(new_donor)['reasons']:
                                st.markdown(f"<div style='text-align: center;'>- {reason}</div>", unsafe_allow_html=True)      

                    
                    # Display registered donors
                    st.markdown("<h2 class='sub-header'>Recent Donors</h2>", unsafe_allow_html=True)
                    recent_donors = st.session_state.donors.tail(5).copy()
                    display_columns = ['id', 'age', 'genre', 'taille', 'poids', 'situation_matrimoniale', 'eligibilite_au_don', 'date_dernier_don', 'taux_dhemoglobine']
                    st.dataframe(recent_donors[display_columns])
        
        if dataset_button:
            with st.form("_form"):
                st.markdown(f"""<div class="card-title" style="text-align: center; font-size: 24px; font-weight: bold; color: #FF5A5A;">
                    {get_text("Donor Registration")}
                    </div> """,  unsafe_allow_html=True)
                col = st.columns(2)
                with col[0] : 
                    Groupe_sanguin= st.selectbox(
                                get_text ("District of Residence"),
                                options=load_data3()["Groupe Sanguin ABO / Rhesus"].unique(),
                                index=0,  # Sélectionner le premier élément par défaut
                                help="Sélectionnez votre groupe"
                            )
                    Age =st.number_input(get_text("Age"), min_value=0, max_value=120, step=1)
                with col[1] : 
                    Horo = st.date_input(get_text("Filling Date"))
                    Genre_ = st.radio(get_text("Gender"), ["Homme", "Femme"])
                rhesus= st.selectbox(
                    get_text ("District of Residence"),
                    options=load_data3()["Phenotype"].unique(),
                    index=0,  # Sélectionner le premier élément par défaut
                    help="Sélectionnez votre Phénotype"
                )
                st.markdown("");st.markdown("");st.markdown("")
                if st.form_submit_button(get_text("Submit_")):
                    if not Age or not rhesus :
                        st.error("Please fill all required fields. It seems that you don't fill the date of birth")
                    else:
                        new_donor = pd.DataFrame({
                        'age': [Age],
                        'genre': [Genre_],
                        'rhesus': [rhesus],
                        'Horo': [Horo.strftime('%Y-%m-%d')],
                        'Groupe_sanguin': [Groupe_sanguin]
                    })
                        st.success(f"Thank you! Donor has been successfully registered.")
                        excel_file = "donnees.xlsx"
                        book = load_workbook(excel_file)

                        # Vérifier si 'Feuil1' existe
                        if 'Feuil2' in book.sheetnames:
                            with pd.ExcelWriter(excel_file, engine='openpyxl', mode='a', if_sheet_exists='overlay') as writer:
                                # Écrire le DataFrame à partir de la première ligne vide
                                startrow = book['Feuil2'].max_row  # Trouver la première ligne vide
                                new_donor.to_excel(writer, index=False, header=False, startrow=startrow, sheet_name='Feuil2')
    


            st.markdown(f"<h2 class='sub-header'>{get_text('NOUVELLE BASE DE DONNEES')}</h2>", unsafe_allow_html=True)
            st.dataframe(pd.read_excel("donnees.xlsx"))

    elif selected_item == "Deep insights":
        st.title("Deep insights")
        st.write("Here you can track Deep insights")
        with st.sidebar:
            donors_button = st.sidebar.button("Donors", use_container_width=True)

    elif selected_item == "Cartography":
        # Initialize session state for button states if they don't exist
        if 'choropleth_active' not in st.session_state:
            st.session_state.choropleth_active = False
        if 'marquer_active' not in st.session_state:
            st.session_state.marquer_active = False
        if 'current_view' not in st.session_state:
            st.session_state.current_view = "none"
        
        # Button functions to toggle states
        def toggle_choropleth():
            st.session_state.choropleth_active = True
            st.session_state.marquer_active = False
            st.session_state.current_view = "choropleth"
        
        def toggle_marquer():
            st.session_state.marquer_active = True
            st.session_state.choropleth_active = False
            st.session_state.current_view = "marquer"
        
        with st.sidebar:
            choropleth_button = st.sidebar.button(get_text("Choropleth"), key="Choropleth", on_click=toggle_choropleth, use_container_width=True)
            marquer_button = st.sidebar.button(get_text("Marquers"), key="Marquers", on_click=toggle_marquer, use_container_width=True)
        
        st.markdown("""
            <style>
            .block-container {
                padding-top: 1.9rem;
                padding-bottom: 0rem;
                padding-left: 1rem;
                padding-right: 2rem;
            }
            </style>
            """, unsafe_allow_html=True)
        

        # Affichage de la carte choroplèthe
        if st.session_state.choropleth_active or not st.session_state.marquer_active:
            gdf = load_shapefile("gadm41_CMR_0.shp")
            m = folium.Map(location=[7.87, 11.52], zoom_start=6, control_scale=True, tiles="CartoDB positron", height=901)
            folium.GeoJson(
                gdf.to_crs(epsg=4326).__geo_interface__,
                name="Cameroun",
                style_function=lambda x: {'fillColor': 'transparent', 'color': 'black', 'weight': 2, 'fillOpacity': 0.0}
            ).add_to(m)
            m.fit_bounds([[4.2, 7.8], [8, 17.75]])
            col1, col2 = st.columns([3.5, 1.5])
            # Créer la carte de base seulement si nécessaire
            with col2:
                st.markdown(f"<h2 class='card-title' style='text-align: center; color: #8a2be2;'>{get_text('Some Filters')}</h2>", unsafe_allow_html=True)
                st.markdown(f"<h2  style='text-align: center; font-size: 16px;'>{get_text('Select a level to watch choropleth map by region, department or district')}</h2>", unsafe_allow_html=True)
                    
                st.markdown("""
                <style>
                div[data-testid="stButton"] button {
                    color: black !important;
                }
                </style>
                """, unsafe_allow_html=True)
            
                if st.button(get_text("By region"), use_container_width=True, key="region_btn")   :
                    # Charger les données prétraitées
                    data_df_1 = get_preprocessed_data(1)
                    st.markdown(f"<h2  style='text-align: center; font-size: 16px;'>{get_text('Top Regions with the highest number of volunteers')}</h2>", unsafe_allow_html=True)
                    st.dataframe(data_df_1, use_container_width=True)
                    
                    # Charger et simplifier les shapefile régionaux
                    geo_data_1 = load_shapefile("gadm41_CMR_1.shp", simplify_tolerance=0.01)
                    geo_data_1 = geo_data_1.merge(data_df_1, on='NAME_1')
                    
                    # Créer une nouvelle carte pour cette vue
                    m = folium.Map(location=[7.87, 11.52], zoom_start=6, control_scale=True, 
                                    tiles="CartoDB positron", height=901)
                    m.fit_bounds([[4.2, 7.8], [8, 17.75]])
                    
                    # Ajouter les contours nationaux
                    gdf = load_shapefile("gadm41_CMR_0.shp", simplify_tolerance=0.01)
                    folium.GeoJson(
                        gdf.to_crs(epsg=4326).__geo_interface__,
                        name="Cameroun",
                        style_function=lambda x: {'fillColor': 'transparent', 'color': 'black', 'weight': 2, 'fillOpacity': 0.0}
                    ).add_to(m)
                    
                    # Ajouter le choropleth
                    Choropleth(
                        geo_data=geo_data_1,
                        data=geo_data_1,
                        columns=['NAME_1', 'Nb'],
                        key_on='feature.properties.NAME_1',
                        fill_color='Reds',
                        fill_opacity=1,
                        line_opacity=0.6,
                    ).add_to(m)
                    
                    # Ajouter les tooltips
                    folium.GeoJson(
                        geo_data_1,
                        style_function=lambda x: {'fillOpacity': 0, 'weight': 0},
                        tooltip=folium.features.GeoJsonTooltip(
                            fields=['NAME_1', 'Nb'],
                            aliases=[get_text('Region:'), 'Nb'],
                            style=("background-color: white; color: #333333; font-family: arial; font-size: 12px; padding: 10px;")
                        )
                    ).add_to(m)
                
                if st.button(get_text("By Department"), use_container_width=True, key="dept_btn"):
                    # Charger les données prétraitées
                    data_df_2 = get_preprocessed_data(2)
                    st.markdown(f"<h2  style='text-align: center; font-size: 16px;'>{get_text('Top Departments with the highest number of volunteers')}</h2>", unsafe_allow_html=True)
                    
                    st.dataframe(data_df_2.head(10), use_container_width=True)
                    
                    # Charger et simplifier les shapefile départementaux
                    geo_data_2 = load_shapefile("gadm41_CMR_2.shp", simplify_tolerance=0.001)
                    geo_data_2 = geo_data_2.merge(data_df_2, on='NAME_2')
                    
                    # Créer une nouvelle carte pour cette vue
                    m = folium.Map(location=[7.87, 11.52], zoom_start=6, control_scale=True, 
                                tiles="CartoDB positron", height=901)
                    m.fit_bounds([[3.2, 8.8], [9, 17.4]])
                    # Ajouter les contours nationaux
                    gdf = load_shapefile("gadm41_CMR_0.shp", simplify_tolerance=0.01)
                    folium.GeoJson(
                        gdf.to_crs(epsg=4326).__geo_interface__,
                        name="Cameroun",
                        style_function=lambda x: {'fillColor': 'transparent', 'color': 'black', 'weight': 2, 'fillOpacity': 0.0}
                    ).add_to(m)
                    
                    # Ajouter le choropleth
                    Choropleth(
                        geo_data=geo_data_2,
                        data=geo_data_2,
                        columns=['NAME_2', 'Nb'],
                        key_on='feature.properties.NAME_2',
                        fill_color='Reds',
                        fill_opacity=1,
                        line_opacity=0.6,
                        highlight=True
                    ).add_to(m)
                    
                    # Ajouter les tooltips
                    folium.GeoJson(
                        geo_data_2,
                        style_function=lambda x: {'fillOpacity': 0, 'weight': 0},
                        tooltip=folium.features.GeoJsonTooltip(
                            fields=['NAME_2','Nb'],
                            aliases=[get_text('Départment:'), 'Nb'],
                            style=("background-color: white; color: #333333; font-family: arial; font-size: 12px; padding: 10px;")
                        )
                    ).add_to(m)    
                if st.button(get_text("By District"), use_container_width=True, key="dist_btn"):
                    # Charger les données prétraitées
                    data_df_3 = get_preprocessed_data(3)
                    st.markdown(f"<h2  style='text-align: center; font-size: 16px;'>{get_text('Top Districts with the highest number of volunteers')}</h2>", unsafe_allow_html=True)
                    st.dataframe(data_df_3.head(100), use_container_width=True)
                    
                    # Charger et simplifier les shapefiles des arrondissements
                    geo_data_3 = load_shapefile("gadm41_CMR_3.shp", simplify_tolerance=0.01)
                    geo_data_3 = geo_data_3.merge(data_df_3, left_on='NAME_3', right_on='Arrondissement_de_résidence_')
                    
                    # Créer une nouvelle carte pour cette vue
                    m = folium.Map(location=[7.87, 11.52], zoom_start=6, control_scale=True, 
                                tiles="CartoDB positron", height=901)
                    m.fit_bounds([[3.2, 8.8], [9, 17.4]])
                    # Ajouter les contours nationaux
                    gdf = load_shapefile("gadm41_CMR_0.shp", simplify_tolerance=0.01)
                    folium.GeoJson(
                        gdf.to_crs(epsg=4326).__geo_interface__,
                        name="Cameroun",
                        style_function=lambda x: {'fillColor': 'transparent', 'color': 'black', 'weight': 2, 'fillOpacity': 0.0}
                    ).add_to(m)
                    
                    # Ajouter le choropleth
                    Choropleth(
                        geo_data=geo_data_3,
                        data=geo_data_3,
                        columns=['NAME_3', 'Nb'],
                        key_on='feature.properties.NAME_3',
                        fill_color='Reds',
                        fill_opacity=1,
                        line_opacity=0.6,
                    ).add_to(m)
                    
                    # Ajouter les tooltips
                    folium.GeoJson(
                        geo_data_3,
                        style_function=lambda x: {'fillOpacity': 0, 'weight': 0},
                        tooltip=folium.features.GeoJsonTooltip(
                            fields=['NAME_3', 'Nb'],
                            aliases=[get_text('District:'), get_text('Volonteers')],
                            style=("background-color: white; color:rgb(15, 15, 15); font-family: arial; font-size: 12px; padding: 10px;")
                        )
                    ).add_to(m)
                    
                
            with col1:
                    with stylable_container(
                        key='1739',
                        css_styles=f"""
                            {{
                                width: {'100%'};   
                                border: 1px solid #c0c0c0;
                                border-radius: 10px;
                                margin: 0px;  # Small margin for spacing
                                padding: 0px;
                                display: flex;
                                flex-direction: column;
                                align-items: center;
                                justify-content: center;
                                background-color: #f8f9fa;
                                box-shadow: 0 4px 6px rgba(0, 0, 0, 0.2);
                                transition: all 0.2s ease;
                                overflow: hidden;
                                text-overflow: ellipsis;
                                color: green;
                            }}
                            
                            .card-title {{
                                    font-weight: bold;
                                    margin: 0px;
                                    padding: 0px;
                                    font-size: 1em;
                                    text-align: center;
                                    color: #8a2be2;  # Light purple color
                                }}
                        """
                    ): folium_static(m, width=900, height=900)
        # Affichage de la carte de marqueurs
        if st.session_state.marquer_active:
            # Charger les données une seule fois
            combined_data = get_hierarchical_data()
            #combined_data = pd.read_excel('Challenge dataset.xlsx')
            #combined_data.rename(columns={'Quartier de_Résidence': 'Arrondissement_de_résidence_'}, inplace=True)
            

            # Charger et simplifier le shapefile une seule fois
            geo_data_3 = load_shapefile("gadm41_CMR_3.shp", simplify_tolerance=0.01)

            # Calculer tous les centroids en une seule fois avec une approche vectorisée
            geo_data_3['centroid_lat'] = geo_data_3['geometry'].centroid.y
            geo_data_3['centroid_lon'] = geo_data_3['geometry'].centroid.x

            # Créer un dictionnaire de centroids plus simplement
            centroid_dict = dict(zip(geo_data_3['NAME_3'], zip(geo_data_3['centroid_lat'], geo_data_3['centroid_lon'])))

            # Fusionner seulement après avoir préparé les centroids
            geo_data_3 = geo_data_3.merge(combined_data, left_on='NAME_3', right_on='Arrondissement_de_résidence_')
            # Identifier les indices des lignes avec coordonnées manquantes
            mask_missing = combined_data['latitude'].isna() | combined_data['longitude'].isna()
            missing_indices = combined_data[mask_missing].index

            # Appliquer les centroids en une seule fois pour chaque colonne
            for idx in missing_indices:
                arrondissement = combined_data.loc[idx, 'Arrondissement_de_résidence_']
                if arrondissement in centroid_dict:
                    combined_data.loc[idx, 'latitude'] = centroid_dict[arrondissement][0]
                    combined_data.loc[idx, 'longitude'] = centroid_dict[arrondissement][1]
    #___________________________________________
            combined_data = combined_data.dropna(subset=['latitude', 'longitude'])
            
            # Créer la carte de base
            if 'marker_map' not in st.session_state or st.session_state.current_view != "marquer":
                gdf = load_shapefile("gadm41_CMR_0.shp", simplify_tolerance=0.01)
                m1 = folium.Map(location=[7.87, 11.52], zoom_start=4, control_scale=True, 
                    tiles="CartoDB positron", height=901)

                folium.GeoJson(
                    gdf.to_crs(epsg=4326).__geo_interface__,
                    name="Cameroun",
                    style_function=lambda x: {'fillColor': 'transparent', 'color': 'black', 'weight': 2, 'fillOpacity': 0.0}
                ).add_to(m1) 
                st.session_state.marker_map = m1
            
            # Colonnes d'affichage
            col1, col2 = st.columns([3.5, 1.5])
            
            with col2:
                st.markdown(f"<h2 class='card-title' style='text-align: center; color: #8a2be2;'>{get_text('Some Filters')}</h2>", unsafe_allow_html=True)
                st.markdown(f"<h7 class='card-title' style='text-align: center; '>{get_text('By default, Nothing selected means all is selected')}</h7>", unsafe_allow_html=True)
                
                all_arrondissement = combined_data['Arrondissement_de_résidence_'].unique()
                selected_arrondissement = st.multiselect(
                    get_text("Districts"),
                    all_arrondissement
                )
                if not selected_arrondissement:
                    selected_arrondissement = all_arrondissement

                # Filtres optimisés
                all_marital = sorted(combined_data['Situation_Matrimoniale_(SM)'].unique())
                selected_marital = st.multiselect(
                    "Marital status",
                    all_marital,  # Limiter par défaut pour améliorer les performances
                )

                if not selected_marital:
                    selected_marital = all_marital

                all_case = combined_data['ÉLIGIBILITÉ_AU_DON.'].unique()
                selected_case = st.multiselect(
                    "Eligible",
                    all_case # Limiter par défaut pour améliorer les performances
                )
                if not selected_case:
                    selected_case = all_case

                all_gender = sorted(combined_data['Genre_'].unique())
                selected_gender = st.multiselect(
                    get_text("Gender"),
                    all_gender # Limiter par défaut pour améliorer les performances
                )
                if not selected_gender:
                    selected_gender = all_gender

                # min_age = int(combined_data['Age'].min())
                # max_age = int(combined_data['Age'].max())
                # age_range = st.slider(
                #     get_text("Select age range"),
                #     min_value=min_age,
                #     max_value=max_age,
                #     value=(min_age, max_age),
                #     step=1
                # )

                all_level = sorted(combined_data["Niveau_d'etude"].unique())
                selected_level = st.multiselect(
                    get_text("Level"),
                    all_level # Limiter par défaut pour améliorer les performances
                )
                if not selected_level:
                    selected_level = all_level
                
                # Filtrer les données avec tous les filtres
                filtered_data = combined_data[
                    (combined_data['Arrondissement_de_résidence_'].isin(selected_arrondissement)) &
                    (combined_data['Situation_Matrimoniale_(SM)'].isin(selected_marital)) &
                    (combined_data['ÉLIGIBILITÉ_AU_DON.'].isin(selected_case)) &
                    (combined_data['Genre_'].isin(selected_gender)) &
                    #(combined_data['Age'] >= age_range[0]) &
                    #(combined_data['Age'] <= age_range[1]) &
                    (combined_data["Niveau_d'etude"].isin(selected_level))
                    
    ]

                # Grouper par quartier pour optimiser le rendu des marqueurs
                quartier_counts = filtered_data['Quartier_de_Résidence_'].value_counts().reset_index()
                quartier_counts.columns = ['Quartier_de_Résidence_', 'count']
                
                # Fusionner avec les coordonnées
                quartier_locations = filtered_data.groupby('Quartier_de_Résidence_').agg({
                    'latitude': 'first',
                    'longitude': 'first'
                }).reset_index()
                
                quartier_data = quartier_counts.merge(quartier_locations, on='Quartier_de_Résidence_')
                
                # Déterminer les tailles des marqueurs
                min_count = quartier_data['count'].min() if not quartier_data.empty else 1
                max_count = quartier_data['count'].max() if not quartier_data.empty else 1
                
                # Créer une nouvelle carte
                m1 = folium.Map(location=[7.87, 11], zoom_start=6, control_scale=True,height=901)

                folium.TileLayer(
                    tiles="CartoDB positron",
                    attr="CartoDB",
                    name="CartoDB Positron",
                    overlay=False,
                    control=True
                ).add_to(m1)
                
                # Ajouter les contours nationaux
                gdf = load_shapefile("gadm41_CMR_0.shp", simplify_tolerance=0.01)
                folium.GeoJson(
                    gdf.to_crs(epsg=4326).__geo_interface__,
                    name="Cameroun",
                    style_function=lambda x: {'fillColor': 'transparent', 'color': 'black', 'weight': 2, 'fillOpacity': 0.0}
                ).add_to(m1)
                m1.fit_bounds([[4.2, 7.8], [8, 17.75]])
                
                # Ajouter les marqueurs de manière optimisée
                for _, row in quartier_data.iterrows():
                    quartier = row["Quartier_de_Résidence_"]
                    count = row["count"]
                    
                    # Calculer la taille du marqueur
                    radius =  3+4* ((count - min_count) / max(max_count - min_count, 1)) * 10
                    
                    folium.CircleMarker(
                        location=[row["latitude"], row["longitude"]],
                        radius=radius,
                        tooltip=f"{quartier}: {count} {get_text('Volonteers')}",
                        color='red',
                        fill=True,
                        fill_color='red'
                    ).add_to(m1)
                
                # Ajouter contrôle de couches
                folium.LayerControl().add_to(m1)
                
            # Afficher la carte
            with col1:
                st.markdown(f"<h3 class='card-title' style='text-align: center; ;'>{get_text('Zoom for better appreciation')}</h3>", unsafe_allow_html=True)
                with stylable_container(
                    key='1946',
                    css_styles=f"""
                        {{
                            width: {'100%'};   
                            border: 1px solid #c0c0c0;
                            border-radius: 10px;
                            margin: 0px;
                            padding: 0px;
                            display: flex;
                            flex-direction: column;
                            align-items: center;
                            justify-content: center;
                            background-color: #f8f9fa;
                            box-shadow: 0 4px 6px rgba(0, 0, 0, 0.2);
                            transition: all 0.2s ease;
                            overflow: hidden;
                            text-overflow: ellipsis;
                            color: green;
                        }}

                    """
                ): folium_static(m1, width=900, height=900)              

    elif selected_item == "About":
        with st.sidebar:
            Blood_Donation = st.sidebar.button(get_text("Blood Donation"), use_container_width=True)
            contact_us  = st.sidebar.button(get_text("Contact Us"), use_container_width=True)

        if Blood_Donation:
            st.markdown(f"<h1 class='main-header'>{get_text('About Blood Donation')}</h1>", unsafe_allow_html=True)
            
            st.markdown(f"""
            <h2 class='sub-header'>{get_text('Why Donate Blood?')}</h2>
            <p class='info-text'>{get_text('Blood donation is a critical lifesaving process that helps millions of people every year. A single donation can save up to three lives, and someone needs blood every two seconds.')}</p>
            
            <div class='highlight'>
            <h3>{get_text('Benefits of Donating Blood:')}</h3>
            <ul>
                <li>{get_text('Helps save lives')}</li>
                <li>{get_text('Free health screening')}</li>
                <li>{get_text('Reduces risk of heart disease')}</li>
                <li>{get_text('Reduces risk of cancer')}</li>
                <li>{get_text('Helps in weight loss')}</li>
                <li>{get_text('Helps in replenishing blood cells')}</li>
            </ul>
            </div>
            """, unsafe_allow_html=True)
            
            col1, col2 = st.columns(2)
            
            with col1:
                st.markdown(f"""
                <h2 class='sub-header'>{get_text('Eligibility Requirements')}</h2>
                <ul>
                    <li>{get_text('Age: 18-65 years')}</li>
                    <li>{get_text('Weight: At least 50 kg')}</li>
                    <li>{get_text('Hemoglobin: 12.5 g/dL for women, 13.0 g/dL for men')}</li>
                    <li>{get_text('Good general health')}</li>
                    <li>{get_text('No fever or active infection')}</li>
                    <li>{get_text('No recent tattoos or piercings (within 4 months)')}</li>
                    <li>{get_text('No recent major surgery')}</li>
                    <li>{get_text('No high-risk behaviors')}</li>
                </ul>
                """, unsafe_allow_html=True)
            
            with col2:
                st.markdown(f"""
                    <h2 class='sub-header'>{get_text('The Donation Process')}</h2>
                    <ol>
                        <li><strong>{get_text('Registration:')}</strong> {get_text('Complete a donor registration form')}</li>
                        <li><strong>{get_text('Health History:')}</strong> {get_text('Answer questions about your health history')}</li>
                        <li><strong>{get_text('Mini-Physical:')}</strong> {get_text('Check temperature, pulse, blood pressure, and hemoglobin')}</li>
                        <li><strong>{get_text('Donation:')}</strong> {get_text('The actual donation takes about 8-10 minutes')}</li>
                        <li><strong>{get_text('Refreshments:')}</strong> {get_text('Rest and enjoy refreshments for 15 minutes')}</li>
                    </ol>
                """, unsafe_allow_html=True)
            
            st.markdown(f"""
            <h2 class='sub-header'>{get_text('Frequently Asked Questions')}</h2>
            """, unsafe_allow_html=True)
            
            faq_expander = st.expander(get_text("Click to view FAQs"))
            
            st.markdown(f"""
            <h2 class='sub-header'>{get_text('Blood Donation Facts')}</h2>
            <div class='highlight'>
            <ul>
                <li>{get_text('Every 2 seconds someone needs blood')}</li>
                <li>{get_text('A single car accident victim can require up to 100 units of blood')}</li>
                <li>{get_text('One donation can save up to 3 lives')}</li>
                <li>{get_text('Only 37% of the population is eligible to donate blood')}</li>
                <li>{get_text('Less than 10% of eligible donors actually donate')}</li>
                <li>{get_text('Blood cannot be manufactured – it can only come from donors')}</li>
                <li>{get_text('Red blood cells have a shelf life of 42 days')}</li>
                <li>{get_text('Platelets have a shelf life of just 5 days')}</li>
            </ul>
            </div>
            """, unsafe_allow_html=True)
        
        if contact_us:
            # About Page Title
            st.markdown(
                f"""
                <h1 style='color: red; text-align: center;'>{get_text('About Our Team')}</h1>
                """,
                unsafe_allow_html=True
            )
            
            # Introduction Text
            st.markdown(
                get_text("Welcome to the About section! Our team is composed of passionate and skilled engineering students in statistics and economics from the Institut Sous-Régional de Statistique et d Économie Appliquée (ISSEA). We are dedicated to building insightful and interactive data-driven solutions, leveraging our expertise in statistical modeling, data visualisations, and economic analysis.") 
            )

            # Display Team Members in a 2x2 Grid
            col1, col2 = st.columns(2)
            @st.cache_resource
            def get_image_path(image_name):
                if image_name.startswith("http"):
                    return image_name  # Pas besoin de modifier les URLs externes
                else:
                    # Si c'est un chemin local, ajuster pour qu'il soit relatif au dossier img
                    # Enlever le "/" initial si présent
                    if image_name.startswith("/"):
                        image_name = image_name[1:]
                    return image_name

            team_members = [
                {"name": "ASSA ALLO", "email": "alloassa21@gmail.com", "phone": "+1237 _________", "image": "img/aa.jpg"},
                {"name": " TAKOUGOUM Steeve Rodrigue", "email": "rodriguetakougoum@gmail.com", "phone": "+237__________", "image": "img/sr.png"},
                {"name": " TIDJANI Razak", "email": "tidjanirazak0@gmail.com", "phone": "+237 ___________", "image": "img/rz.jpg"},            
                {"name": "TCHINDA Chris Donald", "email": "tcd9602@gmail.com", "phone": "+237 ___________ ", "image": "img/tcd.jpg"},
            ]

            # Display Team Members in a 1x4 Grid
            cols = st.columns(4)
            for i, member in enumerate(team_members):
                with cols[i % 4]:
                    image_path = get_image_path(member['image'])
                    
                    st.image(image_path, width=180)
                    
                    st.markdown(
                        f"""
                        <h6>{member['name']}</h6>
                        <p style="font-size: 13px;">✉️ {member['email']}</p>
                        <p style="font-size: 11px;">📞 {member['phone']}</p>
                        """,
                        unsafe_allow_html=True
                    )
                    ""

    elif selected_item=='AI Extensions' :
        import logging
        import subprocess
        import sys
        from dotenv import load_dotenv
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.oxml import OxmlElement
        from datetime import datetime
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_PARAGRAPH_ALIGNMENT

        # Configurer le logging
        logging.basicConfig(level=logging.INFO)
        logger = logging.getLogger(__name__)

        def install(package):
            subprocess.check_call([sys.executable, "-m", "pip", "install", package])

        # Check and install required packages
        required_packages = {
            "transformers": "transformers",
            "groq": "groq",
            "docx": "python-docx",
            "matplotlib": "matplotlib"
        }

        for module, package in required_packages.items():
            try:
                __import__(module)
            except ImportError as e:
                logger.warning(f"{module} package not found: {e}")
                install(package)

        # Import after potential installation
        try:
            from groq import Groq
        except ImportError as e:
            logger.error(f"Failed to import Groq after installation attempt: {e}")
            Groq = None

        # Import document libraries
        try:
            from docx import Document
            from docx.shared import Inches
        except ImportError as e:
            logger.error(f"Failed to import python-docx: {e}")

        load_dotenv()

        # API keys
        groq_api_key = os.getenv("groq_api_key")

        # Function to inject CSS
        def inject_css(css_file):
            try:
                with open(css_file) as f:
                    st.markdown(f'<style>{f.read()}</style>', unsafe_allow_html=True)
            except FileNotFoundError:
                logger.warning(f"CSS file {css_file} not found")

        # Try to inject CSS
        try:
            inject_css('styles.css')
        except Exception as e:
            logger.warning(f"Could not inject CSS: {e}")



        tab_rapport, tab_chatbot_dons, tab_chatbot_medecin = st.tabs([
            "Rapport Dynamique", 
            "Chatbot Analyse Dons", 
            "Chatbot Médecin"
        ])

        # API and model selection
        api_provider = 'Groq'

        # Initialize client variable
        client = None

        if api_provider == 'Groq':
            client = Groq(api_key=groq_api_key)
        

        # Session state initialization
        if "sessions" not in st.session_state:
            st.session_state.sessions = [{"first_query": None, "history": []}]

        if "current_session_index" not in st.session_state:
            st.session_state.current_session_index = 0

        if "editing_query_index" not in st.session_state:
            st.session_state.editing_query_index = None

        if "selected_graphs" not in st.session_state:
            st.session_state.selected_graphs = []

        if "graph_interpretations" not in st.session_state:
            st.session_state.graph_interpretations = {}

        # Generate dashboard data if not present
        if "dashboard_data" not in st.session_state:
            st.session_state.dashboard_data = load_data1()

        # Function for Groq API interaction
        def get_groq_response(messages):
            if client is None:
                return "API client not initialized. Please check your configuration."
            
            try:
                response = client.chat.completions.create(
                    model='Llama3-8b-8192',
                    messages=messages
                )
                return response.choices[0].message.content
            except Exception as e:
                logger.error(f"Error calling Groq API: {e}")
                return f"Erreur lors de la communication avec l'API ou la connexion n'est stable. \n Vous pouvez regénérer le rapport: {str(e)}"

        def interpret_pie_chart(graph_name, data):
            if api_provider == 'Groq' and client:
                # Vérifie si la colonne existe
                if graph_name in data.columns:
                    # Calcul des fréquences relatives
                    distribution = data[graph_name].value_counts(normalize=True).sort_values(ascending=False)
                    total_counts = data[graph_name].value_counts()
                    
                    top_categories = distribution.head(3)
                    formatted_dist = "\n".join([f"- {cat}: {perc*100:.1f}% ({total_counts[cat]} occurrences)" 
                                                for cat, perc in top_categories.items()])
                    
                    prompt = f"""Voici la répartition des catégories pour le graphique en secteur '{graph_name}' :
                    {formatted_dist}
                    
                    Fournis une interprétation professionnelle et concise en 3-5 phrases en français, comme un statisticien spécialisé dans la santé publique au Cameroun. Mentionne les catégories dominantes, les éventuelles disparités et implications potentielles dans le contexte camerounais."""

                    messages = [
                        {"role": "system", "content": "Tu es un expert en analyse de données médicales et de don de sang."},
                        {"role": "user", "content": prompt}
                    ]
                    
                    return get_groq_response(messages)
                else:
                    return f"Pas de données disponibles pour '{graph_name}'."
            else:
                return f"Analyse du graphique en secteur '{graph_name}': la répartition des catégories montre une dominance claire de certaines classes. Une exploration plus poussée pourrait éclairer les implications sociales ou médicales selon le contexte camerounais."

        # Function to create graphs for the report
        def create_graph_image(data, column, title):    
            buf = prepare_frequency_donut_plot(data[column])
            return buf



        # Function to create a Word document with graphs and interpretations
        def create_word_report(selected_graphs, interpretations, data):
            #insert image
            image_path = "indabax_image.jpg"

        # Ouvre l'image et la convertit en bytes
            
            doc = Document()
            
            with open(image_path, 'rb') as file:
                # Lit le contenu de l'image
                image_bytes = file.read()
                
                # Crée un objet BytesIO à partir des bytes
                bytes_io = io.BytesIO(image_bytes)
                
                # Optionnel : pour vérifier que ça fonctionne, on peut repositionner le curseur au début
                bytes_io.seek(0)
        
            doc.add_picture(bytes_io, width=Inches(6))
            
            doc.add_paragraph("\n")
            
            title = doc.add_heading('RAPPORT D\'ANALYSE DES DONS DE SANG', 0)
            # Centrer le titre
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            run = title.runs[0]
            #run.font.color.rgb = RGBColor(255, 0, 0)  # Rouge
            run.font.name = 'Aptos'
            run.font.size = Pt(24)
            run.font.bold = True

            # Ajouter un encadré bleu autour du titre
            paragraph = title._element
            pPr = paragraph.get_or_add_pPr()
            border = OxmlElement('w:pBdr')

            # Définir les bordures (bleues, 2.25 pt) 
            
            # Add date
            rep = doc.add_paragraph(f"Ce rapport a été créé le {datetime.now().strftime('%d/%m/%Y à %H:%M')}")
            doc.add_paragraph("\n")
            rep.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            
            #doc.add_heading('Table des Matières', level=1)
            #for i, graph_name in enumerate(selected_graphs, start=1):
            # doc.add_paragraph(f"Analyse de :  {graph_name}", style='List Bullet')
            #doc.add_page_break()
            
            # Introduction
            doc.add_heading('Introduction', level=1)
            doc.add_paragraph("Le don de sang constitue un acte altruiste et essentiel pour répondre aux besoins médicaux " 
                            "des populations, notamment dans des contextes où les réserves sanguines sont souvent "
                            "insuffisantes, comme au Cameroun. Dans ce cadre, la compréhension des dynamiques de " 
                            "mobilisation, d’éligibilité et de fidélisation des donneurs est cruciale pour optimiser les campagnes " 
                            "de collecte et garantir un approvisionnement constant en sang. Ce rapport présente une analyse " 
                            "approfondie des données issues d’une campagne de don de sang, s’appuyant sur un tableau de bord " 
                            "interactif conçu pour faciliter la prise de décision. À travers une série d’indicateurs clés, de " 
                            "visualisations graphiques et de recommandations stratégiques, nous examinons les profils des "
                            "volontaires, les raisons d’éligibilité ou d’inéligibilité, ainsi que les tendances de participation selon " 
                            "divers critères démographiques et socio-professionnels. L’objectif est double : d’une part, identifier " 
                            "les facteurs qui favorisent ou entravent la participation au don de sang, et d’autre part, proposer " 
                            "des actions concrètes pour améliorer l’efficacité des futures campagnes, tout en renforçant " 
                            "l’engagement des donneurs, notamment les plus récurrents. Ce travail s’inscrit dans une démarche " 
                            "d’amélioration continue des initiatives de santé publique, en mettant l’accent sur l’inclusion, la " 
                            "sensibilisation et l’optimisation des ressources disponibles.\n "   
                            ""       
                            "   Le présent rapport donne une analyse détaillée des données de don de sang, " 
                            "avec des graphiques sélectionnés et leurs interprétations. "
                            "Ces analyses peuvent aider à optimiser les campagnes futures et à améliorer la gestion des stocks.")
            
            # Add each selected graph with its interpretation
            for graph_name in selected_graphs:
                doc.add_heading(f'Analyse: {graph_name}', level=1)
                
                # Create and add the graph image
                img_buf = create_graph_image(data, graph_name, f"Évolution de {graph_name}")
                doc.add_picture(img_buf)
                
                # Add interpretation
                doc.add_heading('Interprétation', level=2)
                doc.add_paragraph(interpretations.get(graph_name, "Pas d'interprétation disponible"))
                
                # Add spacer
                doc.add_paragraph("\n")
            
            # Conclusion
            if selected_graphs:
                doc.add_heading('Conclusion', level=1)
                doc.add_paragraph("L'analyse des graphiques sélectionnés révèle des informations importantes "
                                "sur les tendances des dons de sang. Ces insights peuvent être utilisés pour "
                                "optimiser les campagnes futures et améliorer la gestion des stocks sanguins.")
            
            # Save to BytesIO object
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io

        # Function to create download link
        def get_binary_file_downloader_html(bin_file, file_label='File'):
            bin_str = base64.b64encode(bin_file.read()).decode()
            href = f'<a href="data:application/octet-stream;base64,{bin_str}" download="{file_label}">Télécharger le rapport</a>'
            return href

        # Page 2: Dynamic Report
        with tab_rapport:
            st.title("Rapport Dynamique - Créer et Télécharger")
            
            # Explanation
            st.write("""
            Créez un rapport personnalisé en sélectionnant les graphiques qui vous intéressent.
            Chaque graphique sera accompagné d'une interprétation générée par l'IA.
            Vous pourrez ensuite télécharger le rapport au format Word.
            """)
            
            # Select graphs for the report
            st.subheader("Sélectionner les graphiques pour le rapport")
            
            available_graphs = [col for col in st.session_state.dashboard_data.columns[2:] ]
            selected_graphs = st.multiselect(
                "Choisissez les graphiques à inclure dans le rapport",
                available_graphs,
                default=st.session_state.selected_graphs
            )
            
            # Update the selected graphs in session state
            st.session_state.selected_graphs = selected_graphs
            
            # Display preview of selected graphs
            if selected_graphs:
                st.subheader("Aperçu des graphiques sélectionnés")
                
                # Use columns to display multiple graphs side by side
                for i in range(0, len(selected_graphs), 2):
                    cols = st.columns(2)
                    for j in range(2):
                        if i + j < len(selected_graphs):
                            graph_name = selected_graphs[i + j]
                            with cols[j]:
                                st.write(f"**{graph_name}**")
                                render_frequency_pieh(load_data1()[graph_name], legend_left="70%", legend_top="65%")
                                
                                # Generate or display interpretation
                                if graph_name not in st.session_state.graph_interpretations:
                                    with st.spinner(f"Génération de l'interprétation pour {graph_name}..."):
                                        interpretation = interpret_pie_chart(graph_name, st.session_state.dashboard_data)
                                        st.session_state.graph_interpretations[graph_name] = interpretation
                                
                                st.write("**Interprétation:**")
                                st.write(st.session_state.graph_interpretations[graph_name])
                
                # Option to regenerate interpretations
                if st.button("Régénérer toutes les interprétations"):
                    with st.spinner("Régénération des interprétations..."):
                        st.success("Interprétations régénérées avec succès!")
                        for graph_name in selected_graphs:
                            st.session_state.graph_interpretations[graph_name] = interpret_pie_chart(
                                graph_name, st.session_state.dashboard_data
                            )
                        
                
                # Generate and download report
                st.subheader("Générer et télécharger le rapport")
                
                report_title = st.text_input("Titre du rapport", "Rapport d'Analyse des Dons de Sang")
                
                if st.button("Générer le rapport Word"):
                    with st.spinner("Génération du rapport en cours..."):
                        # Create the Word document
                        report_io = create_word_report(
                            selected_graphs,
                            st.session_state.graph_interpretations,
                            st.session_state.dashboard_data
                        )
                        
                        # Display download link
                        st.success("Rapport généré avec succès!")
                        st.markdown(
                            get_binary_file_downloader_html(
                                report_io, 
                                f"{report_title.replace(' ', '_')}.docx"
                            ),
                            unsafe_allow_html=True
                        )
            else:
                st.info("Veuillez sélectionner au moins un graphique pour créer le rapport.")

        # Page 3: Blood Donation Analysis Chatbot
        with tab_chatbot_dons:
            st.title("Chatbot - Analyse des Dons de Sang")

            if "messages_dons" not in st.session_state:
                st.session_state.messages_dons = [
                    {"role": "assistant", "content": "Salut ! Je peux analyser les données de la campagne de don de sang et te donner des conseils. Pose-moi une question sur les graphiques ou les tendances !"}
                ]

            for msg in st.session_state.messages_dons:
                with st.chat_message(msg["role"]):
                    st.write(msg["content"])

            if prompt := st.chat_input("Pose une question sur les données ou demande un conseil"):
                st.session_state.messages_dons.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.write(prompt)
                    context = f"Données de la campagne de don de sang (résumé): {st.session_state.dashboard_data}. Question: {prompt}"
                
                messages = [
                    {"role": "system", "content": "Tu e expert en analyse de données sur les campagnes de don de sang. Interprète les graphiques; fais des analyses precises, ne fournit aucun code  et donne des conseils contextuels basés sur les données fournies."},
                    {"role": "user", "content": context}
                ]

                with st.chat_message("assistant"):
                    if api_provider == 'Groq':
                        reply = get_groq_response(messages)
                    #else:
                        #reply = get_huggingface_response(context, generator)
                        
                    st.write(reply)
                    st.session_state.messages_dons.append({"role": "assistant", "content": reply})

        # Page 4: Medical Chatbot
        with tab_chatbot_medecin:
            st.title("Chatbot - Médecin Virtuel")

            if "messages_medecin" not in st.session_state:
                st.session_state.messages_medecin = [
                    {"role": "assistant", "content": "Bonjour ! Je suis un médecin virtuel. Pose-moi n'importe quelle question médicale, je te répondrai avec précision et te donnerai des conseils adaptés."}
                ]

            for msg in st.session_state.messages_medecin:
                with st.chat_message(msg["role"]):
                    st.write(msg["content"])

            if prompt := st.chat_input("Pose une question médicale"):
                st.session_state.messages_medecin.append({"role": "user", "content": prompt})
                with st.chat_message("user"):
                    st.write(prompt)

                messages = [
                    {"role": "system", "content": "Tu es un médecin expert avec des années d'expérience. Réponds aux questions médicales de manière précise, professionnelle et compréhensible. Si nécessaire, donne des conseils pratiques ou recommande de consulter un professionnel en personne. Ne fais pas de diagnostics définitifs, mais offre des informations utiles."},
                    {"role": "user", "content": prompt}
                ]

                with st.chat_message("assistant"):
                    if api_provider == 'Groq':
                        reply = get_groq_response(messages)
                    #else:
                        #reply = get_huggingface_response(prompt, generator)
                        
                    st.write(reply)
                    st.session_state.messages_medecin.append({"role": "assistant", "content": reply})



def file_upload_page():
    st.markdown(
        f"""
        <style>
            .block-container {{
                padding-top: 2.5rem;
                padding-left: 25rem;
                padding-right: 25rem;
                position: relative;
                z-index: 1;
            }}
        </style>
        """,
        unsafe_allow_html=True,
    )
    col = st.columns(3)
    with col[1]:
        st.image("OIP.jpeg", width=400)
    st.title("Upload a File")
    uploaded_file = st.file_uploader("Choose a file", type=["csv", "xls", "xlsx"])
    if uploaded_file is not None:
        try:
            if uploaded_file.name.endswith('.csv'):
                df = pd.read_csv(uploaded_file)
            else:
                df = pd.read_excel(uploaded_file)
            # Store the DataFrame in session state but do NOT proceed until validated
            st.session_state["uploaded_data"] = df
            st.session_state["file_uploaded"] = True  # File is uploaded

            st.write("Data Preview: First seven observations")
            st.dataframe(df.head(7))

            # Show the validate button
            if st.button("Validate File"):
                st.session_state["file_validated"] = True  # Mark file as validated

        except Exception as e:
            st.error(f"Error reading the file: {e}")


# Fonction principale pour contrôler le flux de l'application
# Exécuter l'application
if __name__ == "__main__":
    main()
